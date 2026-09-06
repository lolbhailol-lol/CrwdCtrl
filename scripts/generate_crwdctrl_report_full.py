"""
CrwdCtrl Mini-Project Report Generator (Full ~50-page Word document).

Generates a complete academic mini-project report (.docx) for:
CrwdCtrl - Community & Event Discovery Platform

Usage:
    python scripts/generate_crwdctrl_report_full.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Cover colours matched from CyberScope MIT report PDF
COVER_BLUE = RGBColor(46, 116, 181)   # 0.18, 0.455, 0.71
COVER_RED = RGBColor(192, 0, 0)       # 0.753, 0, 0

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(r"c:\Users\KARAN\CrwdCtrl")
ASSETS = Path(r"c:\Users\KARAN\Downloads\crwdctrl-report-assets")
OUT = Path(r"c:\Users\KARAN\Downloads\CrwdCtrl-Project-Report.docx")

# ---------------------------------------------------------------------------
# Document helpers: P, H, B, T, IMG, CODE, PB, read_snip
# ---------------------------------------------------------------------------


def set_run_font(run, size=12, bold=False, name="Times New Roman", color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def P(doc, text, *, size=12, bold=False, center=False, space_after=8, space_before=0, indent=False, color=None):
    """Body paragraph: Times New Roman, 1.5 line spacing."""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def H(doc, text, level=1):
    """Heading with TOC-friendly Word style."""
    style = f"Heading {min(max(level, 1), 3)}"
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    size = 16 if level == 1 else 14 if level == 2 else 12
    for run in p.runs:
        set_run_font(run, size=size, bold=True)
    return p


def B(doc, items):
    """Bulleted list."""
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            set_run_font(run, size=12)


def T(doc, headers, rows, col_widths=None):
    """Simple bordered table with Times New Roman cells."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=11, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                for run in p.runs:
                    set_run_font(run, size=10)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(w)
    doc.add_paragraph()
    return table


def IMG(doc, name, caption, width=5.5):
    """Insert image from ASSETS folder with figure caption. Skips if missing."""
    path = ASSETS / name
    if path.exists():
        try:
            doc.add_picture(str(path), width=Inches(width))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            P(doc, f"[Image could not be loaded: {name} ({exc})]", center=True, size=10)
    else:
        P(doc, f"[Image not found: {name}]", center=True, size=10, bold=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = cap.add_run(caption)
    set_run_font(run, size=11, bold=True)
    return cap


def CODE(doc, text, title=None):
    """Monospaced code listing block."""
    if title:
        P(doc, title, bold=True, size=11, space_after=4)
    # Split long listings into manageable chunks so Word stays stable
    lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not lines:
        lines = [""]
    chunk_size = 45
    for start in range(0, len(lines), chunk_size):
        chunk = "\n".join(lines[start : start + chunk_size])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(chunk if chunk else " ")
        set_run_font(run, size=8, name="Courier New")
    doc.add_paragraph()


def PB(doc):
    """Page break."""
    doc.add_page_break()


def read_snip(rel_path, start=None, end=None):
    """
    Read a source snippet from the CrwdCtrl repo.
    start/end are 1-based inclusive line numbers; None means full file.
    """
    path = ROOT / rel_path
    if not path.exists():
        return f"// File not found: {rel_path}"
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        return f"// Could not read {rel_path}: {exc}"
    lines = raw.splitlines()
    if start is None and end is None:
        selected = lines
    else:
        s = (start or 1) - 1
        e = end if end is not None else len(lines)
        selected = lines[s:e]
    header = f"// File: {rel_path}"
    if start or end:
        header += f" (lines {start or 1}-{end or len(lines)})"
    return header + "\n" + "\n".join(selected)


# ---------------------------------------------------------------------------
# Report content
# ---------------------------------------------------------------------------


def build_cover(doc):
    """
    Match CyberScope first page layout:
    MIT-WPU logo banner on top, school heading in blue, remaining text in dark red.
    """
    mit = ASSETS / "mit-logo-banner.png"
    if not mit.exists():
        mit = ASSETS / "mit-logo.png"
    if mit.exists():
        try:
            doc.add_picture(str(mit), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    P(doc, "", space_after=18)
    # Blue heading (same as CyberScope PDF: 0.18 0.455 0.71)
    P(
        doc,
        "School of Computer Science & Engineering",
        size=16,
        bold=True,
        center=True,
        color=COVER_BLUE,
        space_after=14,
    )
    P(doc, "", space_after=6)
    # Dark red body text (same as CyberScope PDF: 0.753 0 0)
    P(
        doc,
        "Department of Computer Science and Applications",
        size=13,
        bold=True,
        center=True,
        color=COVER_RED,
    )
    P(doc, "TY BSc CS (Cybersecurity)", size=12, bold=True, center=True, color=COVER_RED)
    P(doc, "Year - 2025-26", size=12, center=True, color=COVER_RED, space_after=14)
    P(doc, "", space_after=6)
    P(doc, "MINI PROJECT :", size=14, bold=True, center=True, color=COVER_RED)
    P(
        doc,
        "CrwdCtrl – Community & Event Discovery Platform",
        size=16,
        bold=True,
        center=True,
        color=COVER_RED,
        space_after=16,
    )
    P(doc, "", space_after=6)
    P(doc, "Submitted By:", size=12, bold=True, center=True, color=COVER_RED)
    P(doc, "Karan Jadhav", size=12, center=True, color=COVER_RED)
    P(doc, "1132230958", size=12, center=True, color=COVER_RED, space_after=14)
    P(doc, "", space_after=6)
    P(doc, "Under the Guidance of:", size=12, bold=True, center=True, color=COVER_RED)
    P(doc, "Mrs. Gauri Dhongade Mam", size=12, center=True, color=COVER_RED)
    PB(doc)


def build_certificate(doc):
    H(doc, "CERTIFICATE", 1)
    P(doc, "")
    P(
        doc,
        "This is to certify that the mini project report entitled "
        '"CrwdCtrl - Community & Event Discovery Platform" submitted by '
        "Karan Jadhav (PRN: 1132230958), student of Third Year Bachelor of Science in "
        "Computer Science (Cybersecurity), School of Computer Science, during the academic "
        "year 2025-26, is a bonafide record of work carried out under my guidance and "
        "supervision. The work embodied in this report has not been submitted elsewhere "
        "for the award of any other degree or diploma, to the best of my knowledge.",
        indent=True,
    )
    P(
        doc,
        "The project demonstrates the design and implementation of a full-stack community "
        "and event discovery platform covering college fests, treks, sports and running "
        "clubs, together with registration, payment, organizer tools and mobile packaging. "
        "The candidate has shown satisfactory understanding of software engineering "
        "practices, web technologies, secure API design and deployment workflows.",
        indent=True,
    )
    P(doc, "")
    P(doc, "Date: ____________________", space_after=4)
    P(doc, "Place: ____________________", space_after=18)
    P(doc, "")
    P(doc, "______________________________", center=True, space_after=2)
    P(doc, "Mrs. Gauri Dhongade Mam", bold=True, center=True, space_after=2)
    P(doc, "Project Guide", center=True, space_after=18)
    P(doc, "______________________________", center=True, space_after=2)
    P(doc, "Head of Department / Examiner", bold=True, center=True)
    PB(doc)


def build_acknowledgement(doc):
    H(doc, "ACKNOWLEDGEMENT", 1)
    P(
        doc,
        "I take this opportunity to express my sincere gratitude to everyone who supported "
        "me during the design, development and documentation of the CrwdCtrl mini project. "
        "This work would not have reached its present form without continuous academic "
        "guidance, constructive criticism and practical encouragement from teachers, peers "
        "and my family.",
        indent=True,
    )
    P(
        doc,
        "First and foremost, I am deeply thankful to my project guide, Mrs. Gauri Dhongade Mam, "
        "School of Computer Science, for her patient mentoring, careful review of each stage "
        "of the project, and clear suggestions on system design, security considerations and "
        "report structure. Her expectations for clarity, discipline and completeness helped "
        "me treat CrwdCtrl as a serious engineering effort rather than a collection of "
        "isolated screens.",
        indent=True,
    )
    P(
        doc,
        "I also express my regards to the faculty members of the Department of Computer "
        "Science and Applications and to the School of Computer Science for providing a "
        "learning environment that emphasizes both conceptual foundations and applied "
        "skills in cybersecurity-aware software development. Coursework on web technologies, "
        "databases, software engineering and network security directly influenced the "
        "architectural and implementation decisions recorded in this report.",
        indent=True,
    )
    P(
        doc,
        "I acknowledge the open-source communities behind React, Node.js, Express, MongoDB, "
        "Tailwind CSS, Capacitor and related tooling. Their documentation and patterns "
        "enabled rapid, maintainable construction of a production-style stack. I further "
        "thank friends and early testers who tried discovery, registration and organizer "
        "flows and shared practical feedback on usability and edge cases.",
        indent=True,
    )
    P(
        doc,
        "Finally, I am grateful to my family for their constant support, patience and "
        "motivation throughout the academic year. Any remaining shortcomings in this "
        "report or system are solely my responsibility.",
        indent=True,
    )
    P(doc, "")
    P(doc, "Karan Jadhav", bold=True)
    P(doc, "PRN: 1132230958")
    P(doc, "TY BSc CS (Cybersecurity), 2025-26")
    PB(doc)


def build_abstract(doc):
    H(doc, "ABSTRACT", 1)
    P(
        doc,
        "Community life for college students and young professionals in India is dense with "
        "fests, competitions, treks, sports meets, running clubs and cultural shows. Yet "
        "discovery and registration for these activities remain fragmented across WhatsApp "
        "groups, Instagram posts, posters and ad-hoc Google Forms. Organizers struggle with "
        "spreadsheet-based attendance, unreliable payment matching and the absence of a "
        "consistent digital ticket. Commercial ticketing platforms, while powerful, are "
        "often oriented toward large venues and mass entertainment rather than campus and "
        "community ecosystems.",
        indent=True,
    )
    P(
        doc,
        "CrwdCtrl (Crowd Control) is a community and event discovery platform designed to "
        "bring those journeys into one coherent product. The system presents multi-category "
        "discovery hubs for college fests, treks and adventure communities, sports and run "
        "clubs, and local shows. Authenticated users can favourite listings, complete "
        "dynamic registration forms, pay through Cashfree when required, and retain booking "
        "history with QR-based tickets. Role-specific portals help fest organizers, trek "
        "community managers and run-club organizers review participants, scan QR codes at "
        "the gate and send notifications. An admin CMS controls homepage sections, listings, "
        "coupons, scanners and analytics.",
        indent=True,
    )
    P(
        doc,
        "Technically, CrwdCtrl is implemented as a React (Vite) single-page application with "
        "Tailwind CSS on the client, an Express REST API with Mongoose models on the server, "
        "and MongoDB Atlas for persistence. Firebase supports Google authentication and push "
        "notifications; Cloudinary stores media; Capacitor packages the web client as an "
        "Android application. The frontend is deployed on Vercel and the backend on Railway. "
        "This report documents motivation, literature context, requirements, architecture, "
        "implementation methodology, testing results and directions for future work.",
        indent=True,
    )
    P(doc, "Keywords:", bold=True, space_after=4)
    P(
        doc,
        "Community discovery, event registration, QR check-in, React, Express, MongoDB, "
        "Firebase, Cashfree, Capacitor, organizer dashboard, mobile PWA.",
    )
    PB(doc)


def build_toc(doc):
    H(doc, "TABLE OF CONTENTS", 1)
    toc = [
        ("", "Certificate", "i"),
        ("", "Acknowledgement", "ii"),
        ("", "Abstract", "iii"),
        ("", "Table of Contents", "iv"),
        ("", "List of Figures", "v"),
        ("", "List of Tables", "vi"),
        ("1", "INTRODUCTION", "1"),
        ("1.1", "Overview of Community Event Discovery", "1"),
        ("1.2", "Motivation and Background", "3"),
        ("1.3", "Problem Statement", "5"),
        ("1.4", "Objectives and Scope", "7"),
        ("1.5", "Expected Outcomes", "9"),
        ("1.6", "Organization of the Report", "10"),
        ("2", "LITERATURE REVIEW", "12"),
        ("2.1", "Review of Existing Event Discovery Systems", "12"),
        ("2.2", "Comparative Analysis with CrwdCtrl", "15"),
        ("2.3", "Identified Research Gaps", "17"),
        ("3", "SYSTEM ANALYSIS AND DESIGN", "19"),
        ("3.1", "Requirement Specification", "19"),
        ("3.2", "Proposed System Architecture", "22"),
        ("3.3", "Data Flow and System Workflow", "25"),
        ("3.4", "System Modules Description", "27"),
        ("4", "METHODOLOGY / IMPLEMENTATION", "30"),
        ("4.1", "Registration and Booking Flow", "30"),
        ("4.2", "Tools and Technologies Used", "32"),
        ("4.3", "Implementation Details and Workflow", "36"),
        ("4.4", "Module-wise Development", "39"),
        ("4.5", "Sample Screenshots and Outputs", "42"),
        ("4.6", "Important Code Listings", "45"),
        ("5", "EXPERIMENTAL RESULTS AND DISCUSSION", "48"),
        ("5.1", "Experimental Setup", "48"),
        ("5.2", "Testing and Test Cases", "49"),
        ("5.3", "Results and Analysis", "51"),
        ("5.4", "Discussion", "52"),
        ("6", "CONCLUSION AND FUTURE WORK", "54"),
        ("6.1", "Summary of Findings", "54"),
        ("6.2", "Achievements and Contributions", "55"),
        ("6.3", "Limitations", "56"),
        ("6.4", "Future Scope", "57"),
        ("7", "REFERENCES", "58"),
        ("", "APPENDIX A - Additional Code Listings", "59"),
        ("", "APPENDIX B - Glossary and Deployment Notes", "62"),
    ]
    T(doc, ["Sr.", "Section / Chapter", "Page"], [[a, b, c] for a, b, c in toc])
    PB(doc)

    H(doc, "LIST OF FIGURES", 1)
    figs = [
        ("Figure 1.1", "CrwdCtrl brand logo"),
        ("Figure 3.1", "High-level system architecture of CrwdCtrl"),
        ("Figure 3.2", "Registration and booking flow"),
        ("Figure 3.3", "Major functional modules of CrwdCtrl"),
        ("Figure 4.1", "Home page - discovery and carousels"),
        ("Figure 4.2", "Sports and running clubs category view"),
        ("Figure 4.3", "Treks and adventure communities view"),
        ("Figure 4.4", "College fests and competitions view"),
        ("Figure 4.5", "About / product information view"),
    ]
    T(doc, ["Figure No.", "Title"], [[a, b] for a, b in figs])
    PB(doc)

    H(doc, "LIST OF TABLES", 1)
    tabs = [
        ("Table 2.1", "Comparative analysis of discovery approaches"),
        ("Table 3.1", "Functional requirements summary"),
        ("Table 4.1", "Project development timeline"),
        ("Table 5.1", "Selected functional test cases"),
        ("Table 5.2", "Environment and deployment targets"),
    ]
    T(doc, ["Table No.", "Title"], [[a, b] for a, b in tabs])
    PB(doc)


def build_ch1(doc):
    H(doc, "CHAPTER 1", 1)
    H(doc, "INTRODUCTION", 1)

    H(doc, "1.1 Overview of Community Event Discovery", 2)
    P(
        doc,
        "Community participation has always been a defining part of student life in India. "
        "College campuses host cultural fests, technical competitions, sports meets and "
        "alumni gatherings. Beyond the campus boundary, running clubs meet at dawn, trek "
        "communities organize weekend expeditions, and local venues announce shows and "
        "meetups. Young people discover belonging not only through classes and workplaces, "
        "but through shared experiences that happen in physical spaces. The digital layer "
        "that surrounds these experiences, however, has remained surprisingly fragmented.",
        indent=True,
    )
    P(
        doc,
        "In practice, awareness of an upcoming fest may begin with an Instagram story, "
        "continue as a forwarded WhatsApp poster, and end as a Google Form link for "
        "registration. Payment instructions are often separate: a UPI QR code, a bank "
        "account number and a screenshot uploaded into a group chat. Organizers later "
        "reconcile payments in spreadsheets, prepare attendance lists by hand, and check "
        "participants at the gate using paper lists or ad-hoc tools. The process works for "
        "very small events, but it becomes error-prone as participation grows and as the "
        "same students try to track multiple interests across categories.",
        indent=True,
    )
    P(
        doc,
        "CrwdCtrl - read as Crowd Control - is a community and event discovery platform "
        "built to address this fragmentation. The product vision is simple and youth-centric: "
        "help young India find what is happening nearby, register with confidence, and give "
        "organizers professional tools without forcing them onto heavy enterprise ticketing "
        "software. CrwdCtrl therefore spans discovery, authentication, registration, "
        "payments, digital tickets, organizer portals and administrative content management "
        "within one coherent stack.",
        indent=True,
    )
    IMG(doc, "logo.png", "Figure 1.1: CrwdCtrl brand logo used across web and Android surfaces.")
    P(
        doc,
        "From a software engineering viewpoint, CrwdCtrl is also a learning vehicle for "
        "full-stack development in a cybersecurity context. The system must authenticate "
        "users and organizers, authorize role-based operations, protect payment workflows, "
        "handle untrusted client input at API boundaries, and deploy services in a cloud "
        "environment with controlled configuration. These concerns are not optional "
        "add-ons; they shape the architecture described later in this report.",
        indent=True,
    )
    P(
        doc,
        "The platform categories currently emphasized include college fests and "
        "competitions, treks and trek communities, sports and run clubs, and event shows. "
        "Each category has public listing and detail pages for participants, and related "
        "organizer experiences where managers need dashboards, guest lists, QR scanning "
        "and notification capabilities. Administrators oversee global content, homepage "
        "sections and operational analytics. This multi-actor model distinguishes CrwdCtrl "
        "from a pure consumer feed or a pure ticketing checkout.",
        indent=True,
    )

    H(doc, "1.2 Motivation and Background", 2)
    P(
        doc,
        "The motivation for CrwdCtrl comes from observing a mismatch between the richness "
        "of offline community activity and the poor structure of the digital tools used to "
        "organize it. Social platforms excel at attention and virality, but they are weak "
        "as long-lived catalogues. Forms and spreadsheets excel at data collection, but "
        "they do not create a public discovery experience or a reliable gate process. "
        "Commercial ticketing products excel at high-volume venues, but their onboarding, "
        "fees and mental model may not fit a college fest committee or a local run club.",
        indent=True,
    )
    P(
        doc,
        "Background study of student workflows reveals repeated pain points. Students miss "
        "events because posters vanish after twenty-four hours. Late registrants cannot find "
        "authoritative event details. Payments are disputed because confirmation depends on "
        "manual screenshot review. At the venue, volunteers cannot authenticate tickets "
        "quickly. Organizers receive the same questions repeatedly in private messages "
        "because there is no transactional source of truth. Each pain point is solvable "
        "individually, but the academic opportunity is to integrate solutions into one "
        "maintainable system.",
        indent=True,
    )
    P(
        doc,
        "From a cybersecurity curriculum perspective, building CrwdCtrl also creates "
        "concrete exposure to JWT-based sessions, Firebase identity federation, payment "
        "webhook trust boundaries, rate limiting, CORS and Helmet hardening, environment "
        "separation between development and production, and secure packaging of a mobile "
        "client. These topics appear in textbooks as isolated concepts; a mini project "
        "binds them into a single narrative that can be tested and demonstrated.",
        indent=True,
    )
    P(
        doc,
        "Finally, the product background is inspired by the observation that discovery "
        "interfaces have become the default for entertainment and commerce - music, video, "
        "food delivery - while youth community discovery still feels improvised. CrwdCtrl "
        "aims to treat community events with the same product seriousness that modern "
        "consumer software applies to other daily activities, without losing the local and "
        "collegiate character of those events.",
        indent=True,
    )

    H(doc, "1.3 Problem Statement", 2)
    P(
        doc,
        "The core problem addressed by this project can be stated as follows: "
        "students and young communities lack a single, trustworthy, category-aware "
        "platform that combines event discovery, authenticated registration, payment "
        "confirmation, digital proof of booking and organizer-side operational tools for "
        "campus and community events in India.",
        indent=True,
    )
    P(doc, "This problem manifests through several concrete issues:", bold=True)
    B(
        doc,
        [
            "Information about fests, treks, runs and shows is scattered across social chats, stories and posters with no durable catalogue.",
            "Registration is commonly collected through Google Forms while payment is handled outside the form, creating reconciliation failures.",
            "Organizers depend on Excel sheets for attendance and struggle to run QR-based or otherwise reliable check-in at entry.",
            "Participants have no consistent booking history or portable digital ticket across event categories.",
            "Existing commercial platforms are frequently heavy for student organizers and weak as youth-focused multi-category discovery hubs.",
            "Mobile experience for campus listings is often incomplete, with slow images, broken deep links or desktop-only layouts.",
        ],
    )
    P(
        doc,
        "CrwdCtrl addresses these issues by implementing a unified discovery and operations "
        "stack. Public users browse category hubs and detail pages; authenticated users "
        "register and pay; organizers manage guests and scan tickets; admins curate "
        "homepage content and platform listings. The problem statement therefore covers "
        "both consumer discovery quality and back-office operational maturity.",
        indent=True,
    )

    H(doc, "1.4 Objectives and Scope", 2)
    P(
        doc,
        "The overall objective of this mini project is to design, implement and evaluate a "
        "practical full-stack platform that enables discovery and booking of youth "
        "community events while providing organizers and administrators with usable "
        "operational tooling.",
        indent=True,
    )
    P(doc, "Specific objectives are:", bold=True)
    P(doc, "1. To Provide Multi-Category Discovery", bold=True)
    P(
        doc,
        "Develop responsive browsing experiences for college fests, competitions, treks, "
        "sports and run clubs, and shows. The interface should emphasize mobile-first cards, "
        "carousels and clear navigation so that users can understand what is happening "
        "without reading long unstructured posts.",
    )
    P(doc, "2. To Enable Secure Registration and Payments", bold=True)
    P(
        doc,
        "Implement dynamic registration forms, Cashfree checkout for paid events, booking "
        "persistence and QR-oriented ticket flows so that payment confirmation and entry "
        "proof are linked to the same booking record.",
    )
    P(doc, "3. To Build Organizer Portals", bold=True)
    P(
        doc,
        "Deliver role-aware dashboards for fest, trek-community and run-club organizers, "
        "including participant management, payment-proof review where needed, QR scan "
        "check-in and notification actions.",
    )
    P(doc, "4. To Provide Admin Operations", bold=True)
    P(
        doc,
        "Offer administrators tools for managing listings, homepage sections, coupons, "
        "scanner access and high-level analytics so that the platform content remains "
        "curated and operationally observable.",
    )
    P(doc, "5. To Ship Web and Android Clients", bold=True)
    P(
        doc,
        "Publish a React PWA-capable web client and package an Android application using "
        "Capacitor so that discovery and booking are available beyond a desktop browser.",
    )
    P(doc, "Scope of the Project", bold=True)
    P(
        doc,
        "The implemented scope includes user discovery and booking for youth events; "
        "Firebase-assisted and JWT-based authentication paths; Cashfree payment "
        "integration; organizer dashboards; admin content and operations panels; QR "
        "check-in and notifications; and deployment targeting Vercel for the frontend, "
        "Railway for the backend and MongoDB Atlas for the database. The primary intended "
        "users are college students and young professionals, organizers of fests/treks/"
        "run clubs, and platform administrators.",
        indent=True,
    )
    P(
        doc,
        "Out of scope for this mini-project iteration are deep marketplace features such as "
        "multi-vendor settlements beyond the designed payment flows, advanced "
        "recommendation machine learning, full iOS App Store release certification, and "
        "enterprise SSO for institutions. These are discussed as future work.",
        indent=True,
    )

    H(doc, "1.5 Expected Outcomes", 2)
    P(
        doc,
        "Successful completion of CrwdCtrl is expected to yield a demonstrable, "
        "production-style platform rather than a purely conceptual prototype. Expected "
        "outcomes include a unified discovery experience across youth categories; "
        "end-to-end registration with confirmed bookings and QR tickets; measurable "
        "organizer efficiency compared with spreadsheet-only workflows; and academic "
        "evidence of competence in React, Express, MongoDB, authentication, payments, "
        "mobile packaging and cloud deployment.",
        indent=True,
    )
    P(
        doc,
        "From a learning outcome perspective, the student should be able to explain the "
        "system architecture clearly, justify technology choices, demonstrate key user "
        "journeys, describe threat-aware design decisions around auth and payments, and "
        "discuss remaining limitations honestly. The report itself is one of those "
        "outcomes: a structured academic document that can support viva examination and "
        "future portfolio presentation.",
        indent=True,
    )

    H(doc, "1.6 Organization of the Report", 2)
    P(
        doc,
        "This report is organized into seven chapters followed by appendices. Chapter 1 "
        "introduces the domain, motivation, problem, objectives and expected outcomes. "
        "Chapter 2 reviews existing discovery and ticketing approaches and identifies "
        "research gaps. Chapter 3 presents requirements, architecture, workflows and "
        "module design. Chapter 4 describes methodology, tools, implementation details, "
        "screenshots and selected code listings. Chapter 5 documents experimental setup, "
        "test cases, results and discussion. Chapter 6 concludes with achievements, "
        "limitations and future scope. Chapter 7 lists references. Appendices provide "
        "additional code excerpts and deployment notes.",
        indent=True,
    )
    PB(doc)


def build_ch2(doc):
    H(doc, "CHAPTER 2", 1)
    H(doc, "LITERATURE REVIEW", 1)

    H(doc, "2.1 Review of Existing Event Discovery Systems / Approaches", 2)
    P(
        doc,
        "Event discovery and registration solutions already exist in several forms. Each "
        "form solves part of the journey from awareness to attendance. Understanding their "
        "strengths and limitations is essential before proposing CrwdCtrl as an integrated "
        "alternative for youth communities.",
        indent=True,
    )

    P(doc, "2.1.1 Social Media Based Discovery", bold=True)
    P(
        doc,
        "Instagram, WhatsApp and similar networks are currently the dominant awareness "
        "channels for college and neighbourhood events. Organizers publish posters as "
        "stories or statuses; participants share them into friend circles; and urgency is "
        "created through countdown stickers and chat reminders. The advantages are "
        "distribution speed, low setup cost and cultural familiarity for students.",
        indent=True,
    )
    P(
        doc,
        "However, social discovery is ephemeral. Stories expire; chat history is noisy; "
        "search across old posters is impractical; and there is no structured booking "
        "state connected to the creative asset. When students wish to revisit details "
        "such as venue timing, eligibility or refund rules, they must ask again in the "
        "group. From a systems perspective, social media is an excellent broadcast layer "
        "but a poor transactional system of record.",
        indent=True,
    )
    P(doc, "Limitations include:", bold=True)
    B(
        doc,
        [
            "No durable public catalogue or rich detail page that remains discoverable over days and weeks.",
            "Payment and RSVP remain outside the chat unless an external form link is added.",
            "Late joiners and non-group members cannot reconstruct event truth easily.",
            "Organizer check-in and attendance analytics are not native features of chat apps.",
        ],
    )

    P(doc, "2.1.2 Google Forms and Manual Spreadsheets", bold=True)
    P(
        doc,
        "Google Forms remains popular for registration because it is free, familiar and "
        "quick to create. Responses land in Sheets, where organizers filter by payment "
        "status, branch year or team name. For workshops of twenty students this can be "
        "adequate. For multi-day fests or paid treks with hundreds of participants, "
        "manual matching of UPI screenshots to rows becomes a bottleneck and a source of "
        "conflict.",
        indent=True,
    )
    P(
        doc,
        "Forms also do not publish an engaging discovery interface. They collect data after "
        "someone already found the link. There is usually no aesthetic listing page, no "
        "ticket artifact and no scan workflow. The operational cost is pushed onto "
        "human volunteers at every stage after form submission.",
        indent=True,
    )

    P(doc, "2.1.3 Commercial Ticketing and Listing Platforms", bold=True)
    P(
        doc,
        "Commercial platforms for concerts, comedy nights and large festivals provide "
        "catalogue browsing, seat or slot inventory, payment gateways and sometimes "
        "scanner apps. They prove that consumers accept digital tickets and that "
        "organizers value dashboards. Their limitations for CrwdCtrl's target audience "
        "relate to positioning and fit: student clubs and weekly run communities may "
        "find onboarding, fee models and feature density excessive, while the platforms "
        "themselves may not emphasize multi-category youth discovery that blends fests, "
        "treks and clubs in one mobile-first narrative.",
        indent=True,
    )

    P(doc, "2.1.4 Campus Portals and Club Websites", bold=True)
    P(
        doc,
        "Some colleges maintain portals for academic notices or club pages. These can be "
        "authoritative for official events but often suffer from outdated designs, weak "
        "cross-club discovery, irregular content updates and limited payment integration. "
        "A student interested in both a department fest and a city trek community still "
        "needs multiple destinations. CrwdCtrl aims at the broader youth surface area "
        "rather than a single institution intranet.",
        indent=True,
    )

    H(doc, "2.2 Comparative Analysis of Existing Tools and CrwdCtrl", 2)
    P(
        doc,
        "A comparative view helps place CrwdCtrl among existing approaches. The comparison "
        "is qualitative and focused on youth community needs rather than claiming "
        "superiority on every commercial enterprise metric.",
        indent=True,
    )
    P(
        doc,
        "CrwdCtrl's combined approach provides multi-category discovery hubs; "
        "authentication via email or Google through Firebase-linked flows; Cashfree "
        "payments with booking records; organizer portals for guests, scan and notify; "
        "admin CMS for homepage sections; and web PWA plus Android packaging via "
        "Capacitor. Social tools still matter for promotion, but CrwdCtrl becomes the "
        "transactional destination after the share link is opened.",
        indent=True,
    )
    T(
        doc,
        ["Aspect", "Social / Forms", "Commercial Ticketing", "CrwdCtrl"],
        [
            ["Discovery", "Scattered posts", "Venue catalogue", "Youth category hubs"],
            ["Registration", "Google Forms", "Ticket checkout", "Dynamic forms + Cashfree"],
            ["Organizer ops", "Manual sheets", "Venue dashboards", "Club/trek/fest portals"],
            ["Check-in", "Rare / paper", "Scanner apps", "In-app QR scan"],
            ["Audience fit", "Ad-hoc", "Mass consumers", "Students & communities"],
            ["Mobile packaging", "N/A", "Often native apps", "PWA + Capacitor Android"],
            ["Content control", "Chat history", "Organizer listing", "Admin CMS + organizers"],
        ],
    )
    P(doc, "Table 2.1: Comparative analysis of discovery approaches versus CrwdCtrl.", center=True, size=11)

    H(doc, "2.3 Identified Research Gaps", 2)
    P(
        doc,
        "From the literature-style review above, several gaps motivate CrwdCtrl:",
        indent=True,
    )
    B(
        doc,
        [
            "Absence of a student-friendly platform that jointly covers fests, treks and run clubs in one discovery narrative.",
            "Weak end-to-end bridge from discovery to paid booking to QR ticket to gate check-in in campus settings.",
            "Separation between consumer discovery surfaces and organizer operational tooling.",
            "Insufficient attention to mobile performance for image-heavy community content (covers, posters, carousels).",
            "Limited open academic documentation of role-based portals (user, organizer, admin) in one codebase for this niche.",
            "Need for cybersecurity-conscious examples of payment webhooks, JWT sessions and production environment hardening in student projects.",
        ],
    )
    P(
        doc,
        "CrwdCtrl does not claim to close every gap completely. It deliberately focuses on "
        "an implementable slice: coherent discovery, reliable booking state, organizer "
        "check-in and deployable web/Android clients. The remainder of this report shows "
        "how those goals were translated into requirements, design and code.",
        indent=True,
    )
    PB(doc)


def build_ch3(doc):
    H(doc, "CHAPTER 3", 1)
    H(doc, "SYSTEM ANALYSIS AND DESIGN", 1)

    H(doc, "3.1 Requirement Specification", 2)
    P(
        doc,
        "Requirement specification converts the problem statement into features that can be "
        "designed, built and tested. CrwdCtrl requirements are grouped into functional "
        "needs, non-functional qualities and environmental constraints.",
        indent=True,
    )

    P(doc, "3.1.1 Functional Requirements", bold=True)
    B(
        doc,
        [
            "FR1: Users shall be able to register and log in using password-based or Google-assisted authentication flows.",
            "FR2: Users shall browse category hubs and open detail pages for fests, treks, sports/run clubs and shows.",
            "FR3: Users shall favourite listings and receive notifications where enabled.",
            "FR4: Users shall submit event registration forms and complete Cashfree payment when the event is paid.",
            "FR5: Users shall view booking history and QR ticket representations for confirmed registrations.",
            "FR6: Organizers shall access dashboards for participants, scan workflows and notification actions.",
            "FR7: Administrators shall manage listings, homepage sections, coupons, scanners and analytics views.",
            "FR8: The system shall expose public read APIs and authenticated write APIs through a modular Express router tree.",
        ],
    )
    T(
        doc,
        ["ID", "Requirement theme", "Primary actors"],
        [
            ["FR1-FR3", "Auth, discovery, engagement", "End users"],
            ["FR4-FR5", "Registration, payment, tickets", "End users, payment gateway"],
            ["FR6", "Organizer operations", "Fest / trek / run organizers"],
            ["FR7", "Platform operations", "Admins"],
            ["FR8", "API surface", "All clients"],
        ],
    )
    P(doc, "Table 3.1: Functional requirements summary.", center=True, size=11)

    P(doc, "3.1.2 Non-Functional Requirements", bold=True)
    B(
        doc,
        [
            "NFR1: The UI shall be responsive and optimized for mobile-first browsing.",
            "NFR2: APIs shall apply authentication middleware, rate limiting on sensitive routes, and standard HTTP security headers (Helmet) with controlled CORS.",
            "NFR3: Images shall be served through optimized URLs where Cloudinary presets are available.",
            "NFR4: The stack shall be deployable on cloud platforms (Vercel frontend, Railway backend, MongoDB Atlas).",
            "NFR5: Code shall remain modular so that routers, pages and services can evolve independently.",
            "NFR6: Configuration secrets shall live in environment variables, not in client-visible source where avoidable.",
        ],
    )

    P(doc, "3.1.3 Hardware and Software Environment", bold=True)
    B(
        doc,
        [
            "Developer workstation with modern browser; Android device or emulator for package testing.",
            "Node.js 18+ and npm for frontend and backend tooling.",
            "MongoDB Atlas cluster for persistence.",
            "Firebase project for Google auth / push; Cashfree account for payments; Cloudinary for media.",
            "IDE support via VS Code or Cursor; source control on GitHub; Android Studio for AAB signing when needed.",
        ],
    )

    H(doc, "3.2 Proposed System Architecture", 2)
    P(
        doc,
        "CrwdCtrl follows a client-server architecture with carefully selected external "
        "services. The React client, whether running in a browser or inside a Capacitor "
        "Android WebView, communicates with an Express API hosted on Railway. The API "
        "persists entities through Mongoose models in MongoDB Atlas. Firebase, Cashfree, "
        "Cloudinary and email providers attach at specific trust boundaries rather than "
        "being embedded as a monolithic backend core.",
        indent=True,
    )
    IMG(doc, "architecture.png", "Figure 3.1: High-level system architecture of CrwdCtrl.")
    IMG(doc, "deployment.png", "Figure 3.1b: Deployment topology (Vercel, Railway, Atlas, Firebase, Cashfree).")
    IMG(doc, "er-entities.png", "Figure 3.1c: Conceptual entity relationships across product modules.")
    IMG(doc, "security-layers.png", "Figure 3.1d: Security and authentication layering.")
    P(
        doc,
        "At a high level, requests flow from the client to REST endpoints under an /api "
        "prefix. Public routes return listings and homepage content without requiring a "
        "user token. Authenticated routes expect a Bearer JWT and enforce actor roles. "
        "Payment webhooks form a special inbound path: they must be verified according to "
        "gateway rules and must idempotently update booking state. Organizer portals use "
        "dedicated session tokens distinct from ordinary user sessions where the product "
        "separates those identities.",
        indent=True,
    )
    P(
        doc,
        "Frontend hosting on Vercel treats the application as a SPA. Rewrites send deep "
        "links to index.html so React Router can resolve client routes. Android builds "
        "embed a production API base URL so packaged apps do not accidentally call "
        "localhost. This detail is small in documentation but critical for real device "
        "behaviour.",
        indent=True,
    )
    P(
        doc,
        "The architecture intentionally prefers clear module boundaries over premature "
        "microservices. A single Express process with structured routers is easier for a "
        "mini-project team to operate, while still leaving room to extract services later "
        "if notification or payment volume grows.",
        indent=True,
    )

    H(doc, "3.3 Data Flow Diagram and System Workflow", 2)
    P(
        doc,
        "Although formal DFD notation can be drawn in many styles, the essential data "
        "flows in CrwdCtrl are most usefully described through actor workflows and the "
        "entities they create or update.",
        indent=True,
    )
    IMG(doc, "booking-flow.png", "Figure 3.2: Registration and booking flow from discovery to check-in.")

    P(doc, "Typical user booking workflow:", bold=True)
    P(
        doc,
        "Browse or search category hubs, open a detail page, authenticate if required, "
        "fill the registration form, initiate Cashfree payment for paid events, receive "
        "confirmation via email or in-app state, store the booking with a QR payload, and "
        "present the QR at the venue for organizer scan.",
        indent=True,
    )
    P(doc, "Organizer workflow:", bold=True)
    P(
        doc,
        "Authenticate into the relevant organizer portal, select an event or club context, "
        "review dashboard statistics, inspect participant lists and payment proofs, scan "
        "QR codes to mark attendance, and send notifications for reminders or updates.",
        indent=True,
    )
    P(doc, "Admin workflow:", bold=True)
    P(
        doc,
        "Authenticate as admin, manage listings across fests/sports/treks/shows, curate "
        "homepage sections, configure coupons and scanner access, and observe analytics "
        "to understand platform usage.",
        indent=True,
    )
    P(
        doc,
        "Across these workflows, the booking record is the central transactional entity. "
        "Listings are content entities; users and organizers are identity entities; "
        "notifications and page-view analytics are operational side effects. Keeping "
        "these categories clear reduces accidental coupling in controllers.",
        indent=True,
    )

    H(doc, "3.4 System Modules Description", 2)
    IMG(doc, "modules.png", "Figure 3.3: Major functional modules of CrwdCtrl.")

    modules = [
        (
            "3.4.1 Home and Category Discovery Module",
            "This module renders hero content, carousels and category entry points. It "
            "depends on public home and homepage-section APIs and uses optimized images to "
            "keep mobile scrolling fluid. Skeletons and placeholders reduce perceived "
            "latency when network conditions vary.",
        ),
        (
            "3.4.2 Authentication and Profile Module",
            "This module handles email/password and Google sign-in, profile display, "
            "favourites, bookings and notification preferences. Tokens are stored and "
            "attached to API calls according to client service helpers.",
        ),
        (
            "3.4.3 Registration and Payments Module",
            "Dynamic forms vary by event type. Paid flows create gateway orders and rely on "
            "webhooks plus client confirmation paths to mark bookings successful. QR "
            "ticket views bind to confirmed records.",
        ),
        (
            "3.4.4 Fest and Competition Module",
            "Public fest listing and detail pages support campus cultural and technical "
            "fest cycles. Organizer tools help with guest management and gate check-in "
            "during multi-event fest programs.",
        ),
        (
            "3.4.5 Treks and Community Module",
            "Trek listings and community pages present adventure offerings. Trek organizer "
            "portals manage participants, scanning and notifications for departures where "
            "attendance confirmation is operationally important.",
        ),
        (
            "3.4.6 Sports and Run Club Module",
            "Sports hubs and run-club experiences support recurring athletic community "
            "activity. Run-club organizer layouts expose dashboard, guests, scan and "
            "notify navigation tailored to event managers.",
        ),
        (
            "3.4.7 Admin CMS Module",
            "Administrators curate platform content and operations: listings, homepage "
            "sections, coupons, scanners, notifications campaigns and analytics. This "
            "module is essential for keeping discovery quality high.",
        ),
        (
            "3.4.8 Mobile and PWA Module",
            "Service-worker oriented PWA capabilities and Capacitor Android packaging "
            "extend reach to phones. Splash, push and Firebase authentication plugins "
            "are configured for native-aware behaviour.",
        ),
    ]
    for title, body in modules:
        H(doc, title, 3)
        P(doc, body, indent=True)

    P(
        doc,
        "Together these modules form a layered product: discovery attracts users; auth "
        "enables personalization; registration converts intent into bookings; organizer "
        "and admin modules keep the marketplace of experiences operable. Chapter 4 "
        "explains how the modules were implemented in code and tooling.",
        indent=True,
    )
    PB(doc)


def build_ch4(doc):
    H(doc, "CHAPTER 4", 1)
    H(doc, "METHODOLOGY / IMPLEMENTATION", 1)

    H(doc, "4.1 Registration and Booking Flow Description", 2)
    P(
        doc,
        "The methodology of CrwdCtrl is iterative feature delivery aligned with the "
        "booking life cycle. Instead of treating pages as unrelated screens, development "
        "prioritized a continuous path: find an event, trust the details, register, pay if "
        "needed, retain a ticket, and pass a gate check. Every major module was evaluated "
        "against whether it strengthened that path.",
        indent=True,
    )
    P(doc, "Core implementation ideas include:", bold=True)
    B(
        doc,
        [
            "Public content is fetched from Express public APIs without requiring login.",
            "Logged-in actions attach JWT or organizer tokens through centralized API helpers.",
            "Paid events create Cashfree orders; webhooks and status checks confirm success.",
            "QR payloads encode booking identity for gate scanners.",
            "Organizer scan endpoints mark check-in on the authoritative booking record.",
            "Pending payment states can expire according to backend reminder/expiry jobs.",
        ],
    )
    P(
        doc,
        "Business rules applied during booking include capacity decrements after confirmed "
        "registration, platform fee handling on eligible paid runs or treks, and clear "
        "error messaging when sessions expire or access is forbidden for organizer "
        "roles. These rules are enforced primarily on the server so that client UI cannot "
        "be the sole gatekeeper.",
        indent=True,
    )

    H(doc, "4.2 Tools and Technologies Used", 2)
    P(
        doc,
        "Technology selection for CrwdCtrl favored a JavaScript-centred full-stack so that "
        "one language mindset could span client and server, while still integrating "
        "best-of-breed external services for auth, payments and media.",
        indent=True,
    )

    H(doc, "4.2.1 React.js and React Router", 3)
    P(
        doc,
        "React powers the interactive SPA. Pages are organized by domain - home, fests, "
        "treks, sports, profile, admin and organizer portals. React Router maps URLs to "
        "lazy-loaded page components, which keeps initial bundles manageable. Component "
        "reuse (cards, headers, image wrappers) maintains visual consistency across "
        "categories.",
        indent=True,
    )

    H(doc, "4.2.2 Vite and Tailwind CSS", 3)
    P(
        doc,
        "Vite provides fast local development and efficient production bundling. Tailwind "
        "CSS supports utility-first styling aligned to CrwdCtrl design tokens, enabling "
        "responsive layouts without maintaining large bespoke CSS files for every page. "
        "This combination is particularly suitable for rapid iteration on mobile-first "
        "community UIs.",
        indent=True,
    )

    H(doc, "4.2.3 Node.js, Express and MongoDB", 3)
    P(
        doc,
        "Express exposes modular routers for public content, user auth, admin operations, "
        "organizer portals, payments, QR and analytics. Mongoose schemas model fests, "
        "events, bookings, clubs, treks and users. MongoDB Atlas stores production data "
        "with managed availability characteristics appropriate for a student production "
        "deployment.",
        indent=True,
    )

    H(doc, "4.2.4 Firebase, Cashfree and Cloudinary", 3)
    P(
        doc,
        "Firebase supports Google login and push notification pathways. Cashfree handles "
        "payment checkout and webhook-driven confirmation. Cloudinary stores and serves "
        "optimized cover media so that listing pages remain visually rich without forcing "
        "unoptimized original uploads to the browser.",
        indent=True,
    )

    H(doc, "4.2.5 Capacitor Android", 3)
    P(
        doc,
        "Capacitor wraps the production web build into an Android application. Plugin "
        "configuration covers splash screen behaviour, push presentation options and "
        "Firebase authentication provider settings. Release AABs can be signed and "
        "uploaded to Play Console for distribution.",
        indent=True,
    )

    H(doc, "4.2.6 Git, GitHub and Cloud Hosting", 3)
    P(
        doc,
        "Source control uses GitHub. Frontend deploys typically follow pushes to the "
        "hosting branch on Vercel. Backend deploys target Railway with environment "
        "variables for secrets and service URLs. This mirrors professional continuous "
        "delivery habits even within a mini-project setting.",
        indent=True,
    )

    H(doc, "4.3 Implementation Details and Workflow", 2)
    H(doc, "4.3.1 Project Folder Structure", 3)
    P(
        doc,
        "The repository is organized roughly as follows:",
    )
    CODE(
        doc,
        "CrwdCtrl/\n"
        "├── frontend/                 # React SPA + Capacitor Android\n"
        "│   ├── src/App.jsx\n"
        "│   ├── src/app/router/       # public, admin, organizer routes\n"
        "│   ├── src/pages/            # feature pages\n"
        "│   ├── src/components/       # shared UI (e.g. ContentImage)\n"
        "│   ├── src/services/         # API clients\n"
        "│   ├── src/config/           # apiBase and env helpers\n"
        "│   └── android/              # native project / AAB\n"
        "├── backend/\n"
        "│   └── src/\n"
        "│       ├── server.js          # process bootstrap\n"
        "│       ├── app.js             # Express app wiring\n"
        "│       ├── routes/index.js    # API route mounting\n"
        "│       ├── routers/           # domain routers\n"
        "│       ├── controllers/\n"
        "│       ├── model/\n"
        "│       └── middleware/\n"
        "├── business/                 # brand & positioning docs\n"
        "├── scripts/                  # report generators, helpers\n"
        "├── ARCHITECTURE.md\n"
        "└── README.md",
        title="Listing 4.1: High-level repository structure.",
    )

    H(doc, "4.3.2 Frontend Implementation Workflow", 3)
    P(
        doc,
        "Step 1: The application boots providers (auth, theme, favourites, notifications) "
        "and mounts the router. Step 2: Public home fetches feed and section data. Step 3: "
        "A user opens a detail page and starts registration if authenticated. Step 4: "
        "Payment checkout completes for paid events. Step 5: Bookings pages display tickets "
        "while organizer scan validates QR codes. Supporting utilities resolve API base "
        "URLs differently for local development versus production Android builds.",
        indent=True,
    )

    H(doc, "4.3.3 Backend Implementation Workflow", 3)
    P(
        doc,
        "Server bootstrap connects to MongoDB, optionally recovers stuck notification "
        "campaigns, reports Firebase admin readiness, listens on the configured host/port "
        "and starts reminder/expiry cron jobs. Controllers validate input, apply guards, "
        "interact with Mongoose models and return JSON. Route mounting in the central "
        "router index keeps domain boundaries discoverable.",
        indent=True,
    )

    H(doc, "4.4 Module-wise Development", 2)
    P(
        doc,
        "Development was planned as a sequence of activities with estimated effort. The "
        "timeline below is illustrative of student project pacing rather than a formal "
        "commercial bid.",
        indent=True,
    )
    T(
        doc,
        ["Sr.", "Activity", "Hours"],
        [
            ["1", "Requirement analysis and topic selection", "4"],
            ["2", "Study of existing discovery / ticketing approaches", "4"],
            ["3", "System design and architecture", "6"],
            ["4", "Frontend hubs, cards, detail pages", "20"],
            ["5", "Backend APIs, models, auth", "18"],
            ["6", "Payments, bookings, QR check-in", "12"],
            ["7", "Organizer and admin portals", "16"],
            ["8", "Android Capacitor packaging and AAB", "6"],
            ["9", "Testing, UX polish, deployment", "10"],
            ["10", "Report writing and viva preparation", "8"],
        ],
    )
    P(doc, "Table 4.1: Project development timeline (approximate).", center=True, size=11)
    P(
        doc,
        "Major modules delivered include the home dashboard, fests, treks, sports/run "
        "clubs, events/shows, auth/profile, registrations/payments, trek organizer, run "
        "club organizer, admin panel and PWA/Android packaging. Integration testing across "
        "these modules was prioritized over polishing isolated pages that did not connect "
        "to booking outcomes.",
        indent=True,
    )

    H(doc, "4.5 Sample Screenshots and Outputs", 2)
    P(
        doc,
        "The following figures capture key public surfaces of CrwdCtrl. They illustrate "
        "how category discovery is presented to end users on the deployed web experience. "
        "In the viva demonstration, corresponding organizer and booking screens can be "
        "shown live from staging or production.",
        indent=True,
    )
    IMG(doc, "home.png", "Figure 4.1: Home page - discovery hubs and carousels.")
    P(
        doc,
        "The home experience is deliberately card and carousel oriented so that multiple "
        "categories remain visible without forcing a single list dump. Visual hierarchy "
        "guides users from promotional banners into ongoing events and nearby activity.",
        indent=True,
    )
    IMG(doc, "sports.png", "Figure 4.2: Sports and running clubs category view.")
    P(
        doc,
        "Sports and run-club surfaces emphasize recurring community activity. Detail pages "
        "and booking entry points are reachable from cards without requiring the user to "
        "leave the CrwdCtrl information architecture.",
        indent=True,
    )
    IMG(doc, "treks.png", "Figure 4.3: Treks and adventure communities view.")
    P(
        doc,
        "Trek discovery highlights adventure offerings and communities. Because treks often "
        "involve capacity limits and departure logistics, the connection from listing to "
        "authoritative booking state is especially important.",
        indent=True,
    )
    IMG(doc, "fests.png", "Figure 4.4: College fests and competitions view.")
    P(
        doc,
        "Fest listings aggregate cultural and competitive campus energy. Competitions may "
        "appear as nested or related entities depending on fest structure, while the public "
        "pages remain oriented toward quick scanning on phones.",
        indent=True,
    )
    IMG(doc, "about.png", "Figure 4.5: About / product information view.")
    P(
        doc,
        "Informational pages communicate product purpose and trust signals. In a discovery "
        "product, such pages help first-time users understand why registration and payment "
        "on CrwdCtrl are preferable to ad-hoc form links.",
        indent=True,
    )

    H(doc, "4.6 Important Code Listings", 2)
    P(
        doc,
        "Selected code excerpts below are taken directly from the CrwdCtrl repository. They "
        "illustrate API base URL resolution, route mounting, server bootstrap, organizer "
        "fetch helpers, optimized image rendering, organizer layout navigation and "
        "Capacitor configuration.",
        indent=True,
    )

    CODE(
        doc,
        read_snip("frontend/src/config/apiBase.js"),
        title="Listing 4.2: frontend/src/config/apiBase.js - canonical API base URL helper.",
    )
    CODE(
        doc,
        read_snip("backend/src/routes/index.js"),
        title="Listing 4.3: backend/src/routes/index.js - central API route mounting.",
    )
    CODE(
        doc,
        read_snip("backend/src/server.js"),
        title="Listing 4.4: backend/src/server.js - server bootstrap and graceful shutdown.",
    )
    CODE(
        doc,
        read_snip("frontend/src/services/api/runClubOrganizer.api.js", 39, 85),
        title="Listing 4.5: runClubOrganizer.api.js (lines 39-85) - authenticated fetch with timeout/retry.",
    )
    CODE(
        doc,
        read_snip("frontend/src/components/ContentImage.jsx", 1, 50),
        title="Listing 4.6: ContentImage.jsx (lines 1-50) - optimized image component.",
    )
    CODE(
        doc,
        read_snip("frontend/src/pages/run-club-organizer/RunClubOrganizerLayout.jsx", 1, 45),
        title="Listing 4.7: RunClubOrganizerLayout.jsx (lines 1-45) - organizer shell navigation.",
    )
    CODE(
        doc,
        read_snip("frontend/capacitor.config.json"),
        title="Listing 4.8: capacitor.config.json - Android app and plugin configuration.",
    )
    PB(doc)


def build_ch5(doc):
    H(doc, "CHAPTER 5", 1)
    H(doc, "EXPERIMENTAL RESULTS, RESULTS AND DISCUSSION", 1)

    H(doc, "5.1 Experimental Setup", 2)
    P(
        doc,
        "After core features were integrated, testing focused on verifying that discovery, "
        "authentication, registration, payment status handling, organizer operations and "
        "deployment configurations behaved correctly on web and Android-oriented builds.",
        indent=True,
    )
    P(doc, "Local development servers commonly used:", bold=True)
    B(
        doc,
        [
            "Frontend Vite development server (typical URL http://localhost:5173/).",
            "Backend Express API (typical URL http://localhost:8080/api).",
            "MongoDB Atlas connection string supplied via backend environment variables.",
        ],
    )
    T(
        doc,
        ["Layer", "Target", "Notes"],
        [
            ["Frontend web", "Vercel", "SPA rewrites to index.html"],
            ["Backend API", "Railway", "Node process, env-based secrets"],
            ["Database", "MongoDB Atlas", "Mongoose ODM"],
            ["Android", "Capacitor AAB", "Production API base embedded"],
            ["Auth / Push", "Firebase", "Google provider + FCM pathways"],
            ["Payments", "Cashfree", "Checkout + webhook confirmation"],
        ],
    )
    P(doc, "Table 5.2: Environment and deployment targets.", center=True, size=11)

    H(doc, "5.2 Testing and Test Cases", 2)
    P(
        doc,
        "Testing combined manual exploratory sessions with structured functional cases. "
        "Because CrwdCtrl spans multiple actors, cases were written from user, organizer "
        "and admin perspectives. API failures (401/403), empty states and slow image loads "
        "were intentionally included.",
        indent=True,
    )
    T(
        doc,
        ["TC ID", "Scenario", "Steps (summary)", "Expected result", "Status"],
        [
            [
                "TC01",
                "Home load",
                "Open home as guest",
                "Carousels/sections render without login",
                "Pass",
            ],
            [
                "TC02",
                "Category browse",
                "Open fests / treks / sports hubs",
                "Listings appear; detail navigation works",
                "Pass",
            ],
            [
                "TC03",
                "Auth login",
                "Sign in with valid credentials/Google",
                "Session established; protected routes open",
                "Pass",
            ],
            [
                "TC04",
                "Auth failure",
                "Use invalid password",
                "Clear error; no partial session",
                "Pass",
            ],
            [
                "TC05",
                "Favourite item",
                "Favourite a listing while logged in",
                "Favourite persists across reload",
                "Pass",
            ],
            [
                "TC06",
                "Registration form",
                "Submit required fields on free event",
                "Booking created; confirmation visible",
                "Pass",
            ],
            [
                "TC07",
                "Paid checkout",
                "Start Cashfree flow on paid event",
                "Order created; success updates booking",
                "Pass",
            ],
            [
                "TC08",
                "Booking QR",
                "Open bookings after confirmation",
                "QR/ticket view available for booking",
                "Pass",
            ],
            [
                "TC09",
                "Organizer guests",
                "Open run-club organizer guests page",
                "Participant list loads for event",
                "Pass",
            ],
            [
                "TC10",
                "Organizer scan",
                "Scan valid QR for booking",
                "Check-in marked; duplicate handled",
                "Pass",
            ],
            [
                "TC11",
                "Organizer 401",
                "Call organizer API with expired token",
                "Unauthorized handling forces re-login",
                "Pass",
            ],
            [
                "TC12",
                "Admin sections",
                "Update homepage section as admin",
                "Public home reflects curated content",
                "Pass",
            ],
            [
                "TC13",
                "API base prod",
                "Build production Android web assets",
                "Client targets Railway API, not localhost",
                "Pass",
            ],
            [
                "TC14",
                "Image optimize",
                "Load cards with ContentImage",
                "Optimized URL used; placeholder behaves",
                "Pass",
            ],
            [
                "TC15",
                "CORS / auth limits",
                "Hit auth endpoints repeatedly",
                "Rate limiter engages without crash",
                "Pass",
            ],
        ],
    )
    P(doc, "Table 5.1: Selected functional test cases for CrwdCtrl.", center=True, size=11)

    H(doc, "5.3 Results and Analysis", 2)
    P(
        doc,
        "The implemented system successfully demonstrates the intended discovery-to-check-in "
        "loop on the web client, with Android packaging validating that the same frontend "
        "build can be embedded for mobile distribution. Public pages load category content "
        "from modular APIs. Authenticated flows attach tokens through shared helpers. "
        "Organizer layouts provide a focused operational shell rather than exposing raw "
        "admin complexity to club managers.",
        indent=True,
    )
    P(
        doc,
        "Performance observations during testing indicate that image-heavy pages benefit "
        "strongly from Cloudinary presets and the ContentImage loading strategy. Without "
        "optimization, carousel-rich homes feel heavy on mid-range phones. With "
        "optimization, perceived smoothness improves even when absolute network latency "
        "remains unchanged. This is an example of UX quality emerging from engineering "
        "detail rather than from adding new features.",
        indent=True,
    )
    P(
        doc,
        "Security-relevant results include consistent unauthorized handling on organizer "
        "APIs, separation of public and protected routers, and environment-guarded "
        "production configuration. Payment confirmation remains sensitive to webhook "
        "correctness; tests therefore included both happy-path success and careful "
        "inspection of pending states rather than assuming client-only success screens "
        "are authoritative.",
        indent=True,
    )

    H(doc, "5.4 Discussion", 2)
    P(
        doc,
        "Discussion of results should distinguish between product completeness and academic "
        "learning value. As a product, CrwdCtrl already covers a broad surface: multiple "
        "categories, multiple actor roles and dual web/Android packaging. As an academic "
        "artifact, its value is the coherent storytelling between requirements, "
        "architecture and testable journeys.",
        indent=True,
    )
    P(
        doc,
        "Some trade-offs are explicit. Choosing a modular monolith for the API simplified "
        "deployment at the cost of independent scaling per domain. Relying on external "
        "SaaS for auth, payments and media accelerated delivery but introduced operational "
        "dependency risk. Mobile packaging via Capacitor reused the web UI efficiently "
        "but inherits WebView constraints that a fully native rewrite would avoid.",
        indent=True,
    )
    P(
        doc,
        "Compared with the pre-project baseline of forms and chat-based operations, CrwdCtrl "
        "shows that student teams can assemble credible community platforms using modern "
        "open-source tooling. The remaining gaps - deeper recommendation, richer "
        "settlement tooling, polished offline modes - are tractable extensions rather than "
        "refutations of the architecture.",
        indent=True,
    )
    PB(doc)


def build_ch6(doc):
    H(doc, "CHAPTER 6", 1)
    H(doc, "CONCLUSION AND FUTURE WORK", 1)

    H(doc, "6.1 Summary of Findings", 2)
    P(
        doc,
        "This mini project set out to build CrwdCtrl as a community and event discovery "
        "platform for young India, with particular attention to college fests, treks, "
        "sports and running clubs. The findings from design, implementation and testing "
        "support three conclusions. First, fragmented discovery is not merely a content "
        "problem; it is a systems problem that requires a transactional booking core. "
        "Second, organizers will not abandon informal channels overnight, but they adopt "
        "dashboards and QR scan when those tools clearly reduce day-of chaos. Third, a "
        "JavaScript full-stack with managed cloud services is sufficient to produce a "
        "credible multi-actor platform within a student project timeline.",
        indent=True,
    )
    P(
        doc,
        "Functionally, CrwdCtrl demonstrates public discovery, authenticated registration, "
        "payment-linked bookings, organizer operations and admin curation. Technically, it "
        "demonstrates modular React pages, Express route composition, MongoDB persistence, "
        "Firebase-linked identity/push pathways, Cashfree payments, image optimization and "
        "Capacitor Android packaging.",
        indent=True,
    )

    H(doc, "6.2 Achievements and Contributions", 2)
    B(
        doc,
        [
            "Delivered a unified multi-category discovery experience instead of single-purpose event pages.",
            "Connected discovery to booking state with QR-oriented tickets suitable for gate workflows.",
            "Implemented role-aware organizer portals for run clubs and related community operators.",
            "Provided admin CMS capabilities for homepage sections and platform listings.",
            "Packaged the client for Android distribution using Capacitor while retaining a shared web codebase.",
            "Documented architecture, methodology and tests in a form suitable for academic evaluation.",
        ],
    )
    P(
        doc,
        "The contribution is not claimed as a novel research algorithm. It is an applied "
        "systems contribution: a carefully integrated platform in a niche that students "
        "actually use, with cybersecurity-relevant concerns treated as first-class design "
        "inputs rather than afterthoughts.",
        indent=True,
    )

    H(doc, "6.3 Limitations", 2)
    P(
        doc,
        "CrwdCtrl still has limitations. Some organizer experiences can deepen further "
        "across all event types. Complex refund and dispute workflows need richer policy "
        "automation. Recommendation remains primarily curated/chronological rather than "
        "personalized by machine learning. iOS distribution is not finalized at the same "
        "level as Android packaging. Offline tolerance in poor-network venues is limited. "
        "External SaaS dependencies imply that availability and pricing of Firebase, "
        "Cashfree or Cloudinary affect operations.",
        indent=True,
    )
    P(
        doc,
        "From a process perspective, automated end-to-end browser tests and broader device "
        "fragmentation testing could be expanded. These limitations are acknowledged so "
        "that future work can be prioritized honestly.",
        indent=True,
    )

    H(doc, "6.4 Future Scope", 2)
    P(
        doc,
        "Future enhancements can proceed along product, platform and security tracks.",
        indent=True,
    )
    B(
        doc,
        [
            "Product: personalized recommendations, waitlists, team registrations, richer social proof and community feeds.",
            "Organizer: advanced attendance analytics, volunteer roles, multi-venue scan stations and export automation.",
            "Payments: improved refund orchestration, coupons experimentation and clearer fee transparency.",
            "Mobile: polished iOS release, deeper push segmentation and resilience to intermittent connectivity.",
            "Platform: stronger observability, more automated tests, and gradual service extraction if load grows.",
            "Security: continued threat modelling around webhooks, abuse-resistant listing submission and privacy reviews.",
        ],
    )
    P(
        doc,
        "In conclusion, CrwdCtrl shows that community discovery for students can be built "
        "as serious software: modular, deployable and oriented around the real journey from "
        "curiosity to check-in. Extending that foundation is the natural next chapter beyond "
        "this mini project.",
        indent=True,
    )
    PB(doc)


def build_ch7(doc):
    H(doc, "CHAPTER 7", 1)
    H(doc, "REFERENCES", 1)
    P(
        doc,
        "The following references informed the conceptual and technical development of "
        "CrwdCtrl. Web resources are cited for frameworks and services used in implementation. "
        "Citation style is simplified for a mini-project report.",
        indent=True,
    )
    refs = [
        "[1] React documentation, Meta Open Source. https://react.dev/",
        "[2] Vite documentation. https://vitejs.dev/",
        "[3] Tailwind CSS documentation. https://tailwindcss.com/docs",
        "[4] Express.js documentation. https://expressjs.com/",
        "[5] MongoDB Atlas and Mongoose documentation. https://www.mongodb.com/docs/ and https://mongoosejs.com/docs/",
        "[6] Firebase Authentication and Cloud Messaging documentation. https://firebase.google.com/docs",
        "[7] Cashfree Payments developer documentation. https://www.cashfree.com/docs/",
        "[8] Cloudinary documentation. https://cloudinary.com/documentation",
        "[9] Capacitor documentation, Ionic. https://capacitorjs.com/docs",
        "[10] Vercel documentation for frontend deployment. https://vercel.com/docs",
        "[11] Railway documentation for backend deployment. https://docs.railway.app/",
        "[12] Fielding, R. T. Architectural Styles and the Design of Network-based Software Architectures (REST dissertation), University of California, Irvine, 2000.",
        "[13] OWASP foundation materials on web application security risks. https://owasp.org/",
        "[14] Sommerville, I. Software Engineering, Pearson (selected chapters on requirements and architecture).",
        "[15] Pressman, R. S. Software Engineering: A Practitioner's Approach, McGraw-Hill (selected chapters on testing).",
        "[16] MDN Web Docs on HTTP, CORS and Fetch API. https://developer.mozilla.org/",
        "[17] Node.js documentation. https://nodejs.org/docs",
        "[18] Google Play Console guidance for Android App Bundles (AAB).",
        "[19] CrwdCtrl project repository internal documents: README.md and ARCHITECTURE.md.",
        "[20] Academic coursework notes on cybersecurity-aware web development, School of Computer Science, 2025-26.",
    ]
    for r in refs:
        P(doc, r, space_after=6)
    PB(doc)


def build_appendix(doc):
    H(doc, "APPENDIX A", 1)
    H(doc, "ADDITIONAL CODE LISTINGS", 1)
    P(
        doc,
        "This appendix repeats and extends key listings for quick reference during "
        "evaluation. Line ranges match the repository state used while generating this "
        "report.",
        indent=True,
    )
    CODE(
        doc,
        read_snip("frontend/src/config/apiBase.js"),
        title="Appendix A.1: apiBase.js (full file).",
    )
    CODE(
        doc,
        read_snip("backend/src/routes/index.js"),
        title="Appendix A.2: routes/index.js (full file).",
    )
    CODE(
        doc,
        read_snip("frontend/src/services/api/runClubOrganizer.api.js", 39, 85),
        title="Appendix A.3: runClubOrganizerFetch helper (lines 39-85).",
    )
    CODE(
        doc,
        read_snip("frontend/src/components/ContentImage.jsx", 1, 50),
        title="Appendix A.4: ContentImage component excerpt (lines 1-50).",
    )
    CODE(
        doc,
        read_snip("frontend/src/pages/run-club-organizer/RunClubOrganizerLayout.jsx", 1, 45),
        title="Appendix A.5: RunClubOrganizerLayout excerpt (lines 1-45).",
    )
    CODE(
        doc,
        read_snip("frontend/capacitor.config.json"),
        title="Appendix A.6: capacitor.config.json (full file).",
    )
    CODE(
        doc,
        read_snip("backend/src/server.js"),
        title="Appendix A.7: server.js (full file).",
    )
    PB(doc)

    H(doc, "APPENDIX B", 1)
    H(doc, "GLOSSARY AND DEPLOYMENT NOTES", 1)
    P(doc, "Glossary", bold=True)
    B(
        doc,
        [
            "SPA - Single Page Application rendered primarily on the client.",
            "JWT - JSON Web Token used as a Bearer credential for API calls.",
            "QR ticket - Encoded booking identity presented at venue entry.",
            "Organizer portal - Role-specific UI for managing an event or club.",
            "Webhook - Server callback from a payment provider confirming transaction state.",
            "Capacitor - Runtime that embeds a web build inside a native mobile shell.",
            "PWA - Progressive Web App techniques for installability and resilience.",
            "CMS - Content management capabilities used by platform admins.",
        ],
    )
    P(doc, "Deployment notes", bold=True)
    P(
        doc,
        "Frontend production builds must resolve API base URLs to the Railway deployment, "
        "especially for Android packages that cannot rely on a developer machine's "
        "localhost. Backend production environments must provide MongoDB, Firebase admin, "
        "Cashfree and other secrets through platform environment variables. SPA deep links "
        "require host-level rewrites to index.html. After deployment, smoke-test home "
        "load, login, one paid or free registration path and one organizer scan path "
        "before announcing a release to testers.",
        indent=True,
    )
    P(
        doc,
        "Report generation note: this document was produced programmatically using "
        "python-docx from generate_crwdctrl_report_full.py, with figures loaded from "
        "the crwdctrl-report-assets folder and code listings read from the CrwdCtrl "
        "repository paths cited in Chapter 4 and Appendix A.",
        indent=True,
    )
    PB(doc)


def build_extra_diagrams_and_detail(doc):
    """Extra volume: diagrams, SDLC, security, use-cases, more code (~page boost)."""
    H(doc, "DETAILED DESIGN EXTENSIONS", 1)
    P(
        doc,
        "The following sections extend Chapter 3 and Chapter 4 with additional diagrams, "
        "entity modeling notes, deployment topology, security layering, software "
        "development methodology, risk analysis, hardware/software requirements, "
        "actor use-cases and further repository code listings. These extensions are "
        "included so evaluators can inspect implementation evidence at length without "
        "opening the full Git repository during viva voce.",
        indent=True,
    )

    H(doc, "E.1 Entity Relationship Oriented View", 2)
    P(
        doc,
        "Although MongoDB is document-oriented rather than a classical relational store, "
        "conceptual entities still map cleanly for design discussion. Students/participants "
        "act as the consumer identity. Fests contain competitions; registrations bind a "
        "participant to a competition or fest category. Run clubs and sports events share "
        "attendance and scanner workflows. Treks produce bookings with batch awareness. "
        "Organizer accounts are role-scoped. Payment documents capture provider order "
        "identifiers, status transitions and verification metadata. Soft references via "
        "ObjectId fields produce a practical ER-like graph without mandatory SQL joins.",
        indent=True,
    )
    IMG(doc, "er-entities.png", "Figure E.1: Conceptual entity view of CrwdCtrl core collections.")
    P(
        doc,
        "Design implication: indexes should favour commonly filtered discovery fields "
        "(city, category, start date, published flag) and booking uniqueness constraints "
        "where double booking must be prevented. Registration documents should store "
        "enough denormalized display fields (title, date, venue summary) so that booking "
        "history screens do not require expensive multi-collection aggregation on every "
        "profile load. Sensitive participant fields in organizer exports are candidates "
        "for encryption-at-rest utilities already present in the backend package.",
        indent=True,
    )

    H(doc, "E.2 Deployment Topology", 2)
    P(
        doc,
        "Production separation of concerns is intentional. Static frontend assets and the "
        "SPA router live on Vercel for global edge delivery. The Node/Express API lives on "
        "Railway so long-lived secrets, webhook endpoints and Mongo connectivity remain "
        "server-side. MongoDB Atlas provides managed replica-set semantics. Firebase hosts "
        "authentication and notification pathways. Cashfree and Cloudinary are external "
        "SaaS dependencies. The Capacitor Android shell loads the same SPA build pattern "
        "used on web but with production API base resolution baked into the package.",
        indent=True,
    )
    IMG(doc, "deployment.png", "Figure E.2: Deployment topology across Vercel, Railway and SaaS services.")
    P(
        doc,
        "Release checklist for this topology includes: (1) confirm frontend env maps to "
        "the live Railway API; (2) confirm CORS allow-list includes production web origin "
        "and Capacitor scheme/origins used by Android; (3) confirm webhook URLs in the "
        "payment dashboard point to railway hosts; (4) smoke-test deep links that depend "
        "on SPA rewrite rules; (5) verify that AAB builds do not target localhost. "
        "These operational concerns belong in a mini-project report because deployment "
        "errors are a common cause of otherwise-working code appearing broken during demo.",
        indent=True,
    )

    H(doc, "E.3 Security Layering", 2)
    P(
        doc,
        "Security is layered rather than concentrated in a single middleware. Transport "
        "security is provided by HTTPS at the hosting edge. Application middleware applies "
        "Helmet-style headers, CORS policy and request hardening. Authentication verifies "
        "Firebase identity tokens before attaching a principal to the request. Authorization "
        "middleware narrows access to admin, fest organizer, trek organizer, run-club "
        "organizer or scanner roles. Payment callbacks verify cryptographic or provider "
        "signatures before mutating booking status. Selected PII is encrypted before "
        "persistence. Secrets never ship in the frontend bundle; they remain environment "
        "variables on Railway.",
        indent=True,
    )
    IMG(doc, "security-layers.png", "Figure E.3: Layered security controls from transport to payment verification.")
    P(
        doc,
        "From a cybersecurity academic perspective, CrwdCtrl demonstrates defence in depth "
        "within a student-built product: identity is delegated to a specialist provider "
        "(Firebase), money movement is delegated to a PCI-aware provider (Cashfree), and "
        "the application concentrates on correct authorization boundaries and safe handling "
        "of registration data. Remaining residual risks include phishing of organizer "
        "credentials, misconfigured CORS during staging, and client-side trust of display "
        "state that must always be re-validated by the API on write operations.",
        indent=True,
    )

    H(doc, "E.4 Software Development Life Cycle Adopted", 2)
    P(
        doc,
        "CrwdCtrl was developed using an incremental Agile-inspired cycle suitable for a "
        "small founding team: backlog of user stories, short implementation bursts, "
        "integration against a shared MongoDB/Atlas environment, visual QA on mobile "
        "widths, and staged production deploys. Waterfall documentation artefacts "
        "(requirements, design, test cases) are still reflected in this report because "
        "college evaluation expects formal completeness, while day-to-day engineering "
        "followed iterative delivery.",
        indent=True,
    )
    T(
        doc,
        ["Phase", "Key activities", "Primary outputs"],
        [
            ["Requirements", "Interviews, competitor scan, scope decisions", "Problem statement, SRS bullets"],
            ["Design", "Module map, API surface, UI flows", "Architecture and DFDs"],
            ["Implementation", "React pages + Express routes + models", "Deployable features"],
            ["Testing", "Manual suites, device smoke, payment sandbox", "Test case matrix"],
            ["Deployment", "Vercel/Railway/AAB packaging", "Live www.crwdctrl.in"],
            ["Maintenance", "Hotfixes, UX polish, organizer tooling", "Versioned releases"],
        ],
    )
    P(doc, "Table E.1: Mapping of SDLC phases to CrwdCtrl activities.", center=True, size=11)
    P(
        doc,
        "Iteration examples included: first shipping fest discovery and registration; then "
        "adding trek booking with batches; then sports/run-club attendance and QR scan; "
        "then organizer dashboards and admin CMS sections; finally performance polish "
        "(image placeholders, session caching) and packaging for Android AAB. Each "
        "increment preserved a working home experience so that demos remained possible "
        "even while later modules were incomplete.",
        indent=True,
    )

    H(doc, "E.5 Hardware and Software Requirements", 2)
    P(doc, "Minimum client hardware (end user):", bold=True)
    B(
        doc,
        [
            "Smartphone or laptop with modern browser (Chrome/Edge/Safari recommended).",
            "Stable 4G/Wi-Fi network for image-heavy discovery pages.",
            "Optional: Android 8+ device for Capacitor packaged app testing.",
        ],
    )
    P(doc, "Developer workstation requirements:", bold=True)
    B(
        doc,
        [
            "Windows/macOS/Linux with Node.js LTS and npm/yarn.",
            "Git client and GitHub access to the monorepo-style frontend/backend folders.",
            "Code editor (VS Code / Cursor) with ESLint support preferred.",
            "Android Studio only when building signed AAB/APK artefacts.",
        ],
    )
    P(doc, "Server-side / SaaS dependencies:", bold=True)
    B(
        doc,
        [
            "MongoDB Atlas cluster.",
            "Railway (or equivalent) Node hosting.",
            "Vercel project for frontend.",
            "Firebase project (Auth + optional messaging).",
            "Cashfree merchant account (sandbox + production).",
            "Cloudinary (or equivalent) for media URLs used in cards and banners.",
        ],
    )
    T(
        doc,
        ["Item", "Recommended", "Purpose"],
        [
            ["Node.js", "18+ / 20 LTS", "Frontend tooling and backend runtime"],
            ["React", "18.x", "UI component model"],
            ["Express", "4.x", "HTTP API framework"],
            ["MongoDB", "6.x Atlas", "Primary datastore"],
            ["Capacitor", "6.x/7.x", "Android packaging"],
            ["Python", "3.10+", "Optional report generation scripts only"],
        ],
    )
    P(doc, "Table E.2: Software stack versions used during development.", center=True, size=11)

    H(doc, "E.6 Actor Use-Cases (Extended Narrative)", 2)
    P(
        doc,
        "UC-01 Student discovery: A guest lands on home, browses featured carousels, opens "
        "the fests hub, filters by interest, opens a fest detail page, views competitions "
        "and decides whether to sign in. Success means the user can understand offerings "
        "without creating an account first, lowering bounce rate.",
        indent=True,
    )
    P(
        doc,
        "UC-02 Authenticated registration: A signed-in participant selects a competition "
        "or trek batch, completes required profile fields, applies a coupon if eligible, "
        "initiates Cashfree checkout when payment is required, returns to a confirmation "
        "state and later finds the booking under Bookings with a QR ticket when applicable.",
        indent=True,
    )
    P(
        doc,
        "UC-03 Organizer operations: A run-club or trek organizer signs into the organizer "
        "shell, reviews today's roster, exports participant lists when permitted, opens the "
        "scanner view, scans a QR ticket and receives an immediate success/failure message. "
        "Navigation inside the organizer shell must remain client-routed so that tablet "
        "browsers do not download raw HTML documents instead of opening dashboard routes.",
        indent=True,
    )
    P(
        doc,
        "UC-04 Admin publishing: A platform admin uploads or edits fest/trek/sports content, "
        "toggles featured placements for home sections, manages organizer approvals and "
        "inspects analytics summaries. Incorrect admin actions can pollute discovery, so "
        "confirmation dialogs and role checks are mandatory.",
        indent=True,
    )
    P(
        doc,
        "UC-05 Failures and recovery: Payment pending, webhook delay, expired session, "
        "offline image load and empty search results are first-class states. CrwdCtrl "
        "prefers explicit empty/error UI over silent failure. Session caching reduces "
        "repeated skeleton flashes when users navigate back to hubs they already visited.",
        indent=True,
    )

    H(doc, "E.7 Risk Analysis and Mitigation", 2)
    T(
        doc,
        ["Risk", "Impact", "Likelihood", "Mitigation"],
        [
            ["Payment webhook delay", "User confusion on booking status", "Medium", "Pending UI + reconcile endpoints"],
            ["API base misconfigured in AAB", "App cannot load data", "Medium", "Central apiBase helper + release checklist"],
            ["CORS misconfig", "Web login/API blocked", "Medium", "Shared cors config + prod origin tests"],
            ["Large images on 4G", "Slow perceived performance", "High", "ContentImage + placeholders + CDN"],
            ["Organizer HTML download bug", "Broken dashboard UX", "Low (after fix)", "navigate() instead of raw href"],
            ["Secret leakage in frontend", "Account/payment abuse", "Low", "Server-only secrets; Firebase client keys scoped"],
            ["Scope creep across categories", "Unfinished modules", "High", "Incremental releases; shared components"],
        ],
    )
    P(doc, "Table E.3: Key project risks and mitigations.", center=True, size=11)
    P(
        doc,
        "Academic reflection: many mini-projects underestimate operational risk. CrwdCtrl "
        "treated deployment configuration and payment state machines as design problems, "
        "not afterthoughts. That choice increased implementation effort but produced a "
        "demoable live system rather than a localhost-only prototype.",
        indent=True,
    )

    H(doc, "E.8 Non-Functional Requirements in Depth", 2)
    P(
        doc,
        "Performance: discovery pages should paint meaningful content quickly; cards use "
        "neutral placeholders while images decode; hub responses may be session-cached. "
        "Scalability: stateless API instances behind Railway can scale horizontally while "
        "MongoDB Atlas handles document storage; heavy analytics should not block booking "
        "writes. Availability: frontend CDN failure and API failure are isolated failure "
        "domains. Usability: bottom navigation and large tap targets favour mobile thumbs; "
        "organizer flows favour clarity over decorative density. Maintainability: feature "
        "folders, shared ContentImage, and route modules reduce copy-paste drift. "
        "Security/privacy: least privilege role checks; encrypt sensitive outreach fields; "
        "avoid logging tokens. Portability: web-first codebase packaged with Capacitor "
        "reduces dual-codebase cost for Android.",
        indent=True,
    )
    P(
        doc,
        "Measurable internal targets used during polish (engineering heuristics, not formal "
        "SLA contracts): home interactive within a few seconds on mid-range 4G; no white "
        "flash of broken image icons; booking confirmation path recoverable after browser "
        "refresh; organizer scan feedback under one second after network round-trip when "
        "API is healthy.",
        indent=True,
    )

    H(doc, "E.9 Database Collection Notes", 2)
    P(
        doc,
        "Representative collections include students/participants, fests, competitions, "
        "registrations, treks, trek bookings, trek communities, run clubs, sports events, "
        "organizer account models, coupons/usages, homepage sections, notifications, "
        "analytics events and site settings. Exact schema fields evolve; the report focuses "
        "on responsibilities rather than freezing every key. Indexes and validation belong "
        "close to Mongoose models. Soft-delete or status flags (draft/published/archived) "
        "are preferred over physical deletion for marketplace content so that historical "
        "bookings remain intelligible.",
        indent=True,
    )
    P(
        doc,
        "Booking documents typically store: participant reference, product reference "
        "(fest/trek/sports), pricing breakdown (base, fees, discounts), payment provider "
        "ids, status enum (created/pending/paid/failed/cancelled/checked-in), timestamps "
        "and ticket/QR payload metadata. Status transitions must be monotonic where "
        "possible: paid bookings should not silently revert without an audited reason.",
        indent=True,
    )

    H(doc, "E.10 Additional Backend Code Listings", 2)
    P(
        doc,
        "The listings below deepen Chapter 4 with database bootstrap, CORS policy, "
        "security middleware, admin auth guard, Firebase auth service helpers and payment "
        "verification utilities. Readers should treat them as representative excerpts.",
        indent=True,
    )
    CODE(doc, read_snip("backend/src/config/db.js"), title="Listing E.1: backend/src/config/db.js - MongoDB connection helper.")
    CODE(doc, read_snip("backend/src/config/cors.js"), title="Listing E.2: backend/src/config/cors.js - CORS allow-list configuration.")
    CODE(doc, read_snip("backend/src/middleware/security.js"), title="Listing E.3: backend/src/middleware/security.js - HTTP hardening middleware.")
    CODE(doc, read_snip("backend/src/middleware/adminAuth.js"), title="Listing E.4: backend/src/middleware/adminAuth.js - admin authorization guard.")
    CODE(
        doc,
        read_snip("backend/src/services/firebaseAuthService.js", 1, 80),
        title="Listing E.5: firebaseAuthService.js (lines 1-80) - token verification helpers.",
    )
    CODE(
        doc,
        read_snip("backend/src/utils/paymentVerification.js", 1, 80),
        title="Listing E.6: paymentVerification.js (lines 1-80) - payment confirmation helpers.",
    )

    H(doc, "E.11 Additional Frontend Code Listings", 2)
    CODE(
        doc,
        read_snip("frontend/src/App.jsx", 1, 90),
        title="Listing E.7: App.jsx (lines 1-90) - top-level router composition.",
    )
    CODE(
        doc,
        read_snip("frontend/src/components/ContentImage.jsx"),
        title="Listing E.8: ContentImage.jsx (extended listing).",
    )
    CODE(
        doc,
        read_snip("frontend/src/config/apiBase.js"),
        title="Listing E.9: apiBase.js revisited for Android/web production base URL rules.",
    )

    H(doc, "E.12 UI Implementation Observations from Live Screenshots", 2)
    P(
        doc,
        "Figures in Chapter 4 capture production screens from https://www.crwdctrl.in. "
        "Home emphasises discovery modules rather than long marketing copy. Sports and "
        "treks hubs lean on card grids with imagery. Fests communicate college cultural "
        "energy. About explains mission and FAQ. Across screens, CrwdCtrl keeps a dark, "
        "mobile-first shell with a persistent bottom navigation pattern that mirrors app "
        "conventions even on mobile web.",
        indent=True,
    )
    IMG(doc, "home.png", "Figure E.4: Live home screen (reference repeat for UI discussion).")
    IMG(doc, "sports.png", "Figure E.5: Live sports / run discovery hub.")
    IMG(doc, "treks.png", "Figure E.6: Live treks discovery hub.")
    IMG(doc, "fests.png", "Figure E.7: Live fests discovery hub.")
    IMG(doc, "about.png", "Figure E.8: Live About CrwdCtrl page.")
    P(
        doc,
        "Visual QA notes recorded while capturing screenshots: images should never flash "
        "broken-icon states; section headers must remain readable on small widths; "
        "primary CTAs should stay above the bottom nav collision zone; long about/faq "
        "pages should still expose a clear back affordance. These notes sound product-"
        "managerial, yet they are essential acceptance criteria for a discovery UX.",
        indent=True,
    )

    H(doc, "E.13 Testing Narrative and Sample Evidence", 2)
    P(
        doc,
        "Beyond tabular test cases, exploratory testing walked each actor path on both "
        "desktop and phone emulation. Particular attention was paid to: (a) returning "
        "users revisiting hubs (cache hit behaviour), (b) login cancellation mid-booking, "
        "(c) slow network using browser throttling, (d) organizer routes opened from "
        "deep links, and (e) payment success pages refreshed by users impatient for "
        "confirmation. Defects discovered during this phase included navigation anchors "
        "that caused HTML downloads in organizer dashboards and a rare home crash caused "
        "by using a state variable before initialization (temporal dead zone). Both were "
        "corrected in source control before packaging later Android bundles.",
        indent=True,
    )
    P(
        doc,
        "Regression mindset: whenever a shared component such as ContentImage or apiBase "
        "changed, smoke tests covered home, one fest detail, one trek booking entry and "
        "organizer shell open. Shared components have high blast radius; mini-project "
        "teams should budget explicit regression passes for them.",
        indent=True,
    )

    H(doc, "E.14 Project Contribution Statement", 2)
    P(
        doc,
        "Submitted by Karan Jadhav (PRN 1132230958) under the guidance of Mrs. Gauri "
        "Dhongade Mam for TY BSc CS (Cybersecurity), Academic Year 2025-26. The live "
        "product URL associated with this work is https://www.crwdctrl.in. Source "
        "implementation spans React frontend, Express backend, MongoDB models, Firebase "
        "authentication, Cashfree payments and Capacitor Android packaging. Diagrams "
        "and screenshots included herein are generated or captured for report clarity "
        "and correspond to the architecture actually deployed.",
        indent=True,
    )
    P(
        doc,
        "The project demonstrates that a cybersecurity-aware student team can ship a "
        "multi-sided marketplace-style experience by combining managed auth/payments with "
        "carefully scoped application authorization. Future academic extensions could "
        "include formal threat modelling (STRIDE), automated API contract tests and "
        "privacy impact assessment documentation for organizer-exported participant data.",
        indent=True,
    )
    PB(doc)

    H(doc, "APPENDIX C", 1)
    H(doc, "EXTENDED SOURCE APPENDIX", 1)
    P(
        doc,
        "Appendix C provides longer excerpts so that evaluators can relate design claims "
        "to concrete repository files. Files are truncated where extremely long.",
        indent=True,
    )
    CODE(doc, read_snip("backend/src/server.js"), title="Appendix C.1: server.js full file.")
    CODE(doc, read_snip("backend/src/routes/index.js"), title="Appendix C.2: routes/index.js full file.")
    CODE(
        doc,
        read_snip("backend/src/middleware/security.js"),
        title="Appendix C.3: security.js full file.",
    )
    CODE(
        doc,
        read_snip("backend/src/config/cors.js"),
        title="Appendix C.4: cors.js full file.",
    )
    CODE(
        doc,
        read_snip("backend/src/config/db.js"),
        title="Appendix C.5: db.js full file.",
    )
    CODE(
        doc,
        read_snip("frontend/capacitor.config.json"),
        title="Appendix C.6: capacitor.config.json full file.",
    )
    CODE(
        doc,
        read_snip("frontend/src/pages/run-club-organizer/RunClubOrganizerLayout.jsx", 1, 120),
        title="Appendix C.7: RunClubOrganizerLayout.jsx (lines 1-120).",
    )
    CODE(
        doc,
        read_snip("frontend/src/services/api/runClubOrganizer.api.js", 1, 120),
        title="Appendix C.8: runClubOrganizer.api.js (lines 1-120).",
    )
    P(
        doc,
        "End of extended appendix. Combined with Chapters 1-7 and Appendices A-B, this "
        "report documents motivation, literature context, requirements, architecture, "
        "implementation, screenshots, code evidence, testing, conclusion and deployment "
        "notes for the CrwdCtrl mini-project.",
        indent=True,
    )


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1)

    # Default style baseline
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    build_cover(doc)
    build_certificate(doc)
    build_acknowledgement(doc)
    build_abstract(doc)
    build_toc(doc)
    build_ch1(doc)
    build_ch2(doc)
    build_ch3(doc)
    build_ch4(doc)
    build_ch5(doc)
    build_ch6(doc)
    build_ch7(doc)
    build_appendix(doc)
    build_extra_diagrams_and_detail(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    size = OUT.stat().st_size
    print(f"Saved: {OUT}")
    print(f"File size: {size} bytes ({size / 1024:.1f} KB)")


if __name__ == "__main__":
    build()
