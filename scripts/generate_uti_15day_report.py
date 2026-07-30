from pathlib import Path
import json
from datetime import datetime

from docx import Document


ROOT = Path(r"c:\Users\KARAN\CrwdCtrl")
DATA_PATH = ROOT / "reports" / "UTI-15-Day-Report-data.json"
OUT_DOCX = ROOT / "reports" / "UTI-15-Day-Report.docx"
OUT_MD = ROOT / "reports" / "UTI-15-Day-Report.md"


def add_heading(doc, text):
    doc.add_heading(text, level=2)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def format_change(change_percent):
    if change_percent > 0:
        return f"+{change_percent}%"
    if change_percent < 0:
        return f"{change_percent}%"
    return "0%"


def growth_wording(change_percent):
    if change_percent > 0:
        return "increased"
    if change_percent < 0:
        return "decreased"
    return "held steady"


def normalize_device(device):
    label = str(device or "unknown").strip().lower()
    if label == "mobile":
        return "Mobile"
    if label == "desktop":
        return "Desktop"
    return device.title() if device else "Unknown"


def build_source_views(traffic_sources):
    source_views = {"Direct": 0, "Instagram": 0, "Google": 0, "Facebook": 0, "Other": 0}
    for item in traffic_sources:
        src = str(item.get("source", "")).strip().lower()
        hits = int(item.get("hits", 0))
        if src == "direct":
            source_views["Direct"] += hits
        elif src == "instagram":
            source_views["Instagram"] += hits
        elif src == "google":
            source_views["Google"] += hits
        elif src == "facebook":
            source_views["Facebook"] += hits
        else:
            source_views["Other"] += hits
    return source_views


def build_highlights(data, internal, ga4, treks, growth, highlights):
    g = ga4.get("totals", {}) if ga4.get("configured") else {}
    total_trek_views = sum(int(t.get("detailViews", 0)) for t in treks)
    change = int(growth.get("changePercent", 0))
    first7 = growth.get("first7Days", {})
    last7 = growth.get("last7Days", {})
    rank = int(highlights.get("communityRank", 0))
    total_ranked = int(highlights.get("totalCommunitiesRanked", 0))
    top_social = highlights.get("topSocialSource", "Instagram")
    published = int(highlights.get("publishedTreks", len(treks)))
    best_day = highlights.get("bestDay", "—")
    best_day_views = int(highlights.get("bestDayViews", 0))
    engagement = g.get("engagementRate", internal.get("engagementRate", 0))
    bounce = g.get("bounceRate", max(0, 100 - float(engagement or 0)))

    bullets = [
        (
            f"Community page views {growth_wording(change)} week-over-week — "
            f"{first7.get('views', 0)} views in the first 7 days, "
            f"{last7.get('views', 0)} in the last 7 days ({format_change(change)})."
        ),
        f"Best day of the period: {best_day} with {best_day_views} page views.",
        (
            f"Strong on-page engagement — {engagement}% engagement rate and only "
            f"{bounce}% bounce on the community page (Google Analytics)."
        ),
    ]

    if rank == 1 and total_ranked:
        bullets.append(
            f"Ranked #1 among all trek communities on CrwdCtrl for community-page views in this period."
        )
    elif rank and total_ranked:
        bullets.append(
            f"Ranked #{rank} among all trek communities on CrwdCtrl for community-page views in this period."
        )

    bullets.append(
        f"{published} treks live on the platform, driving {total_trek_views} trek page views in 15 days."
    )
    bullets.append(
        f"{top_social} is the top social source — a good signal that social posts are bringing people in."
    )
    return bullets


def write_markdown(content):
    OUT_MD.write_text(content, encoding="utf-8")
    print(f"Wrote {OUT_MD}")


def main():
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    internal = data["internal"]
    ga4 = data.get("ga4", {})
    treks = data.get("treks", [])
    period = data["period"]
    community = data["community"]
    growth = data.get("growth", {})
    highlights = data.get("highlights", {})

    if not growth:
        day_by_day = internal.get("dayByDay", [])
        first7 = day_by_day[:7]
        last7 = day_by_day[-7:]
        growth = {
            "first7Days": {
                "startLabel": first7[0]["label"] if first7 else "",
                "endLabel": first7[-1]["label"] if first7 else "",
                "views": sum(int(d.get("views", 0)) for d in first7),
            },
            "last7Days": {
                "startLabel": last7[0]["label"] if last7 else "",
                "endLabel": last7[-1]["label"] if last7 else "",
                "views": sum(int(d.get("views", 0)) for d in last7),
            },
            "changePercent": 0,
        }
        first_views = growth["first7Days"]["views"]
        last_views = growth["last7Days"]["views"]
        growth["changePercent"] = (
            round(((last_views - first_views) / first_views) * 100)
            if first_views
            else 0
        )

    if not highlights:
        best_day = max(internal.get("dayByDay", []), key=lambda d: int(d.get("views", 0)), default={})
        highlights = {
            "bestDay": best_day.get("label", "—"),
            "bestDayViews": int(best_day.get("views", 0)),
            "publishedTreks": len(treks),
            "topSocialSource": "Instagram",
        }

    total_trek_detail = sum(int(t.get("detailViews", 0)) for t in treks)
    total_trek_book_clicks = sum(int(t.get("bookNowClicks", 0)) for t in treks)
    g = ga4.get("totals", {}) if ga4.get("configured") else {}
    source_views = build_source_views(internal.get("trafficSources", []))
    bullet_lines = build_highlights(data, internal, ga4, treks, growth, highlights)

    start = datetime.fromisoformat(period["startDate"]).strftime("%d %b %Y")
    end = datetime.fromisoformat(period["endDate"]).strftime("%d %b %Y")
    start_short = datetime.fromisoformat(period["startDate"]).strftime("%d %b").lstrip("0")
    end_short = datetime.fromisoformat(period["endDate"]).strftime("%d %b").lstrip("0")
    first7 = growth.get("first7Days", {})
    last7 = growth.get("last7Days", {})

    doc = Document()
    doc.add_heading(f"{community['name']} — Performance Report", level=1)
    doc.add_paragraph(f"Period: {start} – {end} (last {period['days']} days)")
    doc.add_paragraph("Platform: CrwdCtrl")

    doc.add_heading("Highlights", level=2)
    for line in bullet_lines:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "1. Overview")
    add_table(
        doc,
        ["Metric", "Count"],
        [
            ["Active users", g.get("activeUsers", 0)],
            ["Community page views", internal["communityPageViews"]],
            ["Events", g.get("eventCount", 0)],
            ["Trek page views", total_trek_detail],
            ["Clicks to book", total_trek_book_clicks],
            ["Published treks live", len(treks)],
        ],
    )

    add_heading(doc, "2. Growth — first 7 days vs last 7 days")
    add_table(
        doc,
        ["Period", "Community page views"],
        [
            [
                f"First 7 days ({first7.get('startLabel', '')}–{first7.get('endLabel', '')})",
                first7.get("views", 0),
            ],
            [
                f"Last 7 days ({last7.get('startLabel', '')}–{last7.get('endLabel', '')})",
                last7.get("views", 0),
            ],
            ["Change", format_change(int(growth.get("changePercent", 0)))],
        ],
    )

    add_heading(doc, "4. Where traffic came from")
    add_table(doc, ["Source", "Views"], [[k, v] for k, v in source_views.items()])

    add_heading(doc, "5. Device")
    add_table(
        doc,
        ["Device", "Views"],
        [[normalize_device(d.get("device")), d.get("views", 0)] for d in internal.get("deviceSplit", [])],
    )

    add_heading(doc, "6. Trek page views")
    add_table(
        doc,
        ["Trek", "Views", "Clicks to book"],
        [[t["name"], t["detailViews"], t.get("bookNowClicks", 0)] for t in treks]
        + [["Total", total_trek_detail, total_trek_book_clicks]],
    )

    doc.add_paragraph(
        f"Data from CrwdCtrl analytics and Google Analytics 4 · last 15 days ({start_short}–{end_short} {period['endDate'][:4]})."
    )

    md_lines = [
        f"# {community['name']} — Performance Report",
        "",
        f"**Period:** {start} – {end} (last {period['days']} days)  ",
        "**Platform:** CrwdCtrl",
        "",
        "## Highlights",
        "",
    ]
    md_lines.extend(f"- {line}" for line in bullet_lines)
    md_lines.extend(
        [
            "",
            "## 1. Overview",
            "",
            "| Metric | Count |",
            "|---|---|",
            f"| Active users | {g.get('activeUsers', 0)} |",
            f"| Community page views | {internal['communityPageViews']} |",
            f"| Events | {g.get('eventCount', 0)} |",
            f"| Trek page views | {total_trek_detail} |",
            f"| Clicks to book | {total_trek_book_clicks} |",
            f"| Published treks live | {len(treks)} |",
            "",
            "## 2. Growth — first 7 days vs last 7 days",
            "",
            "| Period | Community page views |",
            "|---|---:|",
            f"| First 7 days ({first7.get('startLabel', '')}–{first7.get('endLabel', '')}) | {first7.get('views', 0)} |",
            f"| Last 7 days ({last7.get('startLabel', '')}–{last7.get('endLabel', '')}) | {last7.get('views', 0)} |",
            f"| Change | {format_change(int(growth.get('changePercent', 0)))} |",
            "",
            "## 4. Where traffic came from",
            "",
            "| Source | Views |",
            "|---|---:|",
        ]
    )
    md_lines.extend(f"| {k} | {v} |" for k, v in source_views.items())
    md_lines.extend(
        [
            "",
            "## 5. Device",
            "",
            "| Device | Views |",
            "|---|---:|",
        ]
    )
    md_lines.extend(
        f"| {normalize_device(d.get('device'))} | {d.get('views', 0)} |"
        for d in internal.get("deviceSplit", [])
    )
    md_lines.extend(
        [
            "",
            "## 6. Trek page views",
            "",
            "| Trek | Views | Clicks to book |",
            "|---|---:|---:|",
        ]
    )
    md_lines.extend(
        f"| {t['name']} | {t['detailViews']} | {t.get('bookNowClicks', 0)} |" for t in treks
    )
    md_lines.extend(
        [
            f"| Total | {total_trek_detail} | {total_trek_book_clicks} |",
            "",
            f"Data from CrwdCtrl analytics and Google Analytics 4 · last 15 days ({start_short}–{end_short} {period['endDate'][:4]}).",
            "",
        ]
    )
    write_markdown("\n".join(md_lines))

    try:
        doc.save(OUT_DOCX)
        print(f"Wrote {OUT_DOCX}")
    except PermissionError:
        version = 2
        while True:
            alt_out = OUT_DOCX.with_name(f"{OUT_DOCX.stem}-v{version}{OUT_DOCX.suffix}")
            if not alt_out.exists():
                doc.save(alt_out)
                print(f"Wrote {alt_out}")
                break
            version += 1


if __name__ == "__main__":
    main()
