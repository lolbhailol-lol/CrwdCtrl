"""Generate CrwdCtrl mini-project report (.docx) matching CyberScope report structure."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(r"c:\Users\KARAN\Downloads\CrwdCtrl-Project-Report.docx")


def set_run_font(run, size=12, bold=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, *, size=12, bold=False, center=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_custom(doc, text, level=1):
    # Use paragraph styles for TOC-friendly headings
    style = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in p.runs:
            set_run_font(run, size=12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=11, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # -------- COVER --------
    for _ in range(2):
        add_para(doc, "", size=12)
    add_para(doc, "School of Computer Science & Engineering", size=14, bold=True, center=True)
    add_para(doc, "", size=10)
    add_para(doc, "Department of Computer Science and Applications", size=13, bold=True, center=True)
    add_para(doc, "TY BSc CS (Cybersecurity)", size=12, bold=True, center=True)
    add_para(doc, "Year - 2025-26", size=12, center=True)
    add_para(doc, "", size=12)
    add_para(doc, "MINI PROJECT :", size=14, bold=True, center=True)
    add_para(
        doc,
        "CrwdCtrl – Community & Event Discovery Platform",
        size=16,
        bold=True,
        center=True,
        space_after=12,
    )
    add_para(doc, "", size=12)
    add_para(doc, "Submitted By:", size=12, bold=True, center=True)
    add_para(doc, "Karan Jadhav", size=12, center=True)
    add_para(doc, "1132230958", size=12, center=True)
    add_para(doc, "", size=12)
    add_para(doc, "Under the Guidance of:", size=12, bold=True, center=True)
    add_para(doc, "Mrs. Gauri Dhongade Mam", size=12, center=True)
    page_break(doc)

    # -------- TOC --------
    add_heading_custom(doc, "Table Of Contents", 1)
    toc = [
        ("1", "INTRODUCTION", "3"),
        ("", "1.1 Overview of Community Event Discovery", "3"),
        ("", "1.2 Problem Statement", "4"),
        ("", "1.3 Objectives and Scope", "4"),
        ("", "1.4 Expected Outcomes", "5"),
        ("", "1.5 Organization of the Report", "6"),
        ("2", "LITERATURE REVIEW", "8"),
        ("", "2.1 Review of Existing Event Discovery Systems / Approaches", "8"),
        ("", "2.2 Comparative Analysis of Existing Tools and CrwdCtrl", "10"),
        ("", "2.3 Identified Research Gaps", "11"),
        ("3", "SYSTEM ANALYSIS AND DESIGN", "13"),
        ("", "3.1 Requirement Specification", "13"),
        ("", "3.2 Proposed System Architecture", "15"),
        ("", "3.3 Data Flow Diagram and System Workflow", "17"),
        ("", "3.4 System Modules Description", "18"),
        ("4", "METHODOLOGY / IMPLEMENTATION", "21"),
        ("", "4.1 Registration & Booking Flow Description", "21"),
        ("", "4.2 Tools and Technologies Used", "23"),
        ("", "4.3 Implementation Details and Workflow", "29"),
        ("", "4.4 Module-wise Development", "32"),
        ("", "4.5 Sample Screenshots / Outputs", "40"),
        ("5", "EXPERIMENTAL RESULTS, RESULTS AND DISCUSSION", "45"),
        ("", "5.1 Experimental Setup", "45"),
        ("", "5.2 Testing and Test Cases", "46"),
        ("", "5.3 Results and Analysis", "47"),
        ("", "5.4 Discussion", "48"),
        ("6", "CONCLUSION AND FUTURE WORK", "50"),
        ("", "6.1 Summary of Findings", "50"),
        ("", "6.2 Achievements and Contributions", "51"),
        ("", "6.3 Limitations and Improvements", "52"),
        ("", "6.4 Future Scope", "53"),
        ("7", "REFERENCES", "54"),
    ]
    add_table(doc, ["Sr. No.", "Section", "Page No."], [[a, b, c] for a, b, c in toc])
    page_break(doc)

    # -------- CH 1 --------
    add_heading_custom(doc, "INTRODUCTION", 1)
    add_heading_custom(doc, "1.1 Overview of Community Event Discovery", 2)
    add_para(
        doc,
        "Nowadays college life and youth culture in India are full of fests, competitions, treks, "
        "running clubs, sports meets and cultural shows. Students discover most of these activities "
        "through WhatsApp groups, Instagram stories, posters and Google Forms. Information is "
        "scattered across many places, so people miss events or struggle to register on time.",
    )
    add_para(
        doc,
        "CrwdCtrl (Crowd Control) is a community and event discovery platform built so that young "
        "India can find what is happening nearby in one application. The platform covers college "
        "fests, competitions, treks and adventure communities, sports and running clubs, and "
        "shows / meetups.",
    )
    add_para(doc, "Main idea of CrwdCtrl:", bold=True)
    add_bullets(
        doc,
        [
            "Discover events across categories in one place",
            "Register and pay securely without messy spreadsheets",
            "Give organizers dashboards for participants, QR check-in and notifications",
            "Give admins tools to manage content, sections and analytics",
        ],
    )
    add_para(
        doc,
        "The project brand line used in product messaging is: “Where young India finds what’s happening.” "
        "CrwdCtrl aims to make real-world experiences as easy to discover as streaming a movie or "
        "ordering food.",
    )
    add_para(
        doc,
        "CrwdCtrl – Community & Event Discovery Platform is implemented as a full-stack web "
        "application with an Android app (Capacitor). Users can browse, favourite, register and "
        "receive QR tickets. Organizers manage fests, treks or run clubs. Admins control homepage "
        "sections, listings and operational tools.",
    )

    add_heading_custom(doc, "1.2 Problem Statement", 2)
    add_para(
        doc,
        "Event discovery for students and young communities in India is broken across many informal channels. "
        "A typical student must check multiple Instagram pages, college WhatsApp groups and shared Google Forms "
        "just to know what is happening this weekend.",
    )
    add_para(doc, "Common problems include:", bold=True)
    add_bullets(
        doc,
        [
            "No single trusted catalogue for fests, treks, runs and local shows",
            "Registration done on Google Forms / Excel with payment handled separately (UPI screenshots)",
            "Organizers find it hard to manage attendees, check-in and reminders",
            "No consistent booking history or digital QR ticket for participants",
            "Cold or incomplete listings and poor mobile experience on many campus sites",
        ],
    )
    add_para(
        doc,
        "Existing ticketing products are often built for large venues and concerts. College fests, "
        "community treks and run clubs need simpler discovery plus organizer tools without heavy "
        "enterprise setup. CrwdCtrl solves this by combining discovery, registration, payments "
        "(Cashfree), organizer portals and admin CMS in one stack.",
    )

    add_heading_custom(doc, "1.3 Objectives and Scope", 2)
    add_para(doc, "Objectives", bold=True)
    add_para(
        doc,
        "The main objective of CrwdCtrl is to build a practical full-stack platform that helps users "
        "discover and join youth events, and helps organizers operate those events professionally.",
    )
    add_para(doc, "The objectives of this project are:", bold=True)
    add_para(doc, "1. To Provide Multi-Category Discovery", bold=True)
    add_para(
        doc,
        "Browse and search college fests, competitions, treks, sports / run clubs and events with "
        "mobile-first cards, carousels and category hubs.",
    )
    add_para(doc, "2. To Enable Secure Registration and Payments", bold=True)
    add_para(
        doc,
        "Support dynamic registration forms, Cashfree checkout, booking history and QR-based tickets.",
    )
    add_para(doc, "3. To Build Organizer Portals", bold=True)
    add_para(
        doc,
        "Fest, trek-community and run-club organizers can manage listings, participants, QR scan "
        "check-in, payment proof review and notifications.",
    )
    add_para(doc, "4. To Provide Admin Operations", bold=True)
    add_para(
        doc,
        "Admins can manage fests, sports, treks, homepage sections, coupons, scanners and analytics.",
    )
    add_para(doc, "5. To Ship Web + Android", bold=True)
    add_para(
        doc,
        "Deliver a React PWA on Vercel and a Capacitor Android application packaged as a Play Store AAB.",
    )
    add_para(doc, "Scope of Project", bold=True)
    add_para(doc, "The scope of CrwdCtrl includes:")
    add_bullets(
        doc,
        [
            "User discovery and booking for youth events",
            "Firebase + JWT authentication",
            "Cashfree payment integration",
            "Organizer dashboards (fest / trek / run club)",
            "Admin content and operations panel",
            "QR check-in and notifications",
            "Deployment on Vercel (frontend) and Railway (backend) with MongoDB Atlas",
        ],
    )
    add_para(doc, "The application is made for:")
    add_bullets(
        doc,
        [
            "College students and young professionals",
            "Fest / trek / run-club organizers",
            "Platform administrators",
            "Campus and community event programs",
        ],
    )

    add_heading_custom(doc, "1.4 Expected Outcomes", 2)
    add_para(
        doc,
        "After successful implementation, CrwdCtrl is expected to give a production-style platform "
        "for event discovery and operations.",
    )
    add_para(doc, "Expected outcomes include:", bold=True)
    add_para(doc, "Unified Event Discovery", bold=True)
    add_para(
        doc,
        "Users can find fests, treks, sports / runs and shows from one homepage and category pages.",
    )
    add_para(doc, "End-to-End Registration", bold=True)
    add_para(
        doc,
        "Participants complete forms, pay when required, and receive confirmation / QR ticket flows.",
    )
    add_para(doc, "Organizer Efficiency", bold=True)
    add_para(
        doc,
        "Organizers stop relying only on spreadsheets — they get dashboards for guests, scan and notify.",
    )
    add_para(doc, "Practical Full-Stack Learning", bold=True)
    add_para(
        doc,
        "The project demonstrates React, Express, MongoDB, Firebase auth, Cashfree payments, "
        "Capacitor Android packaging and cloud deployment.",
    )

    add_heading_custom(doc, "1.5 Organization of the Report", 2)
    add_para(doc, "This report is organized into six chapters.")
    add_para(doc, "Chapter 1: Introduction", bold=True)
    add_para(
        doc,
        "This chapter gives an overview of the problem domain, CrwdCtrl concept, problem statement, "
        "objectives, scope and expected outcomes.",
    )
    add_para(doc, "Chapter 2: Literature Review", bold=True)
    add_para(
        doc,
        "This chapter reviews existing event discovery / ticketing approaches, comparison with CrwdCtrl, "
        "and research gaps.",
    )
    add_para(doc, "Chapter 3: System Analysis and Design", bold=True)
    add_para(
        doc,
        "This chapter explains requirements, architecture, workflow and system modules.",
    )
    add_para(doc, "Chapter 4: Methodology and Implementation", bold=True)
    add_para(
        doc,
        "This chapter describes booking flow, technologies, folder structure and module-wise development.",
    )
    add_para(doc, "Chapter 5: Experimental Results and Discussion", bold=True)
    add_para(
        doc,
        "This chapter includes testing environment, sample test cases, results and discussion.",
    )
    add_para(doc, "Chapter 6: Conclusion and Future Work", bold=True)
    add_para(
        doc,
        "This chapter summarizes outcomes, achievements, limitations and future improvements.",
    )
    page_break(doc)

    # -------- CH 2 --------
    add_heading_custom(doc, "LITERATURE REVIEW", 1)
    add_heading_custom(doc, "2.1 Review of Existing Event Discovery Systems / Approaches", 2)
    add_para(
        doc,
        "Many systems already help people find or book events. Most focus on one slice: big concerts, "
        "club parties, or simple Google Form collections. CrwdCtrl combines discovery across campus "
        "and community categories with organizer operations.",
    )
    add_para(doc, "1. Social Media Based Discovery", bold=True)
    add_para(
        doc,
        "Instagram and WhatsApp are the most common discovery layer for college fests and club runs. "
        "Posts and stories spread awareness quickly but information expires fast and is hard to search later.",
    )
    add_para(doc, "Limitations:")
    add_bullets(
        doc,
        [
            "No structured catalogue or booking history",
            "Payment and RSVP stay outside the chat",
            "Hard for late joiners to find old event details",
        ],
    )
    add_para(doc, "2. Google Forms + Manual Spreadsheets", bold=True)
    add_para(
        doc,
        "Many organizers still collect registrations using Google Forms and track payments in Excel. "
        "This works for small events but breaks at scale for check-in and refunds.",
    )
    add_para(doc, "Limitations:")
    add_bullets(
        doc,
        [
            "Manual payment matching (UPI screenshots)",
            "No QR check-in system",
            "No unified public discovery feed",
        ],
    )
    add_para(doc, "3. Commercial Ticketing Platforms", bold=True)
    add_para(
        doc,
        "Products such as district-level ticketing sites focus on paid venues and large festivals. "
        "They are powerful but often heavy for student clubs and weekly run communities.",
    )
    add_para(doc, "Limitations:")
    add_bullets(
        doc,
        [
            "Not optimized for multi-category youth discovery (fests + treks + runs together)",
            "Fee structures and onboarding may not fit campus organizers",
            "Community features (clubs, trek communities) may be missing",
        ],
    )

    add_heading_custom(doc, "2.2 Comparative Analysis of Existing Tools and CrwdCtrl", 2)
    add_para(
        doc,
        "After studying existing approaches, most tools solve only one part of the journey — either "
        "awareness (social) or registration (forms) or tickets (commercial platforms).",
    )
    add_para(doc, "CrwdCtrl gives a combined approach by providing:")
    add_bullets(
        doc,
        [
            "Multi-category discovery (fests, competitions, treks, sports/run clubs, events)",
            "Auth (email/password + Google via Firebase)",
            "Cashfree payments and booking records",
            "Organizer portals with participants, QR scan and notify",
            "Admin CMS for homepage sections and listings",
            "Web PWA + Android Capacitor packaging",
        ],
    )
    add_table(
        doc,
        ["Aspect", "Social / Forms", "Commercial ticketing", "CrwdCtrl"],
        [
            ["Discovery", "Scattered posts", "Events catalogue", "Youth category hubs"],
            ["Registration", "Google Forms", "Ticket checkout", "Dynamic forms + Cashfree"],
            ["Organizer ops", "Manual sheets", "Venue tools", "Club/trek/fest dashboards"],
            ["Check-in", "Rare", "Scanner apps", "In-app QR scan"],
            ["Audience", "Ad-hoc", "Mass consumers", "Students & communities"],
        ],
    )

    add_heading_custom(doc, "2.3 Identified Research Gaps", 2)
    add_para(doc, "Identified gaps that CrwdCtrl addresses:")
    add_bullets(
        doc,
        [
            "No single student-friendly platform combining fests + treks + run clubs",
            "Weak bridge from discovery → paid booking → QR ticket → gate check-in",
            "Organizer tools usually separate from consumer discovery apps",
            "Limited focus on mobile UX and smooth image / carousel performance",
            "Need for role-based portals (user / organizer / admin) in one codebase",
        ],
    )
    page_break(doc)

    # -------- CH 3 --------
    add_heading_custom(doc, "SYSTEM ANALYSIS AND DESIGN", 1)
    add_heading_custom(doc, "3.1 Requirement Specification", 2)
    add_para(
        doc,
        "Requirement specification lists functional features, non-functional needs, and environment "
        "required to run CrwdCtrl.",
    )
    add_para(doc, "Functional Requirements", bold=True)
    add_bullets(
        doc,
        [
            "User registration / login (JWT + Firebase Google)",
            "Browse category hubs and event detail pages",
            "Favourites and notifications",
            "Event registration forms and Cashfree payment",
            "Bookings page with QR ticket views",
            "Organizer dashboards for participants, scan and notify",
            "Admin management for listings, sections and analytics",
        ],
    )
    add_para(doc, "Non-Functional Requirements", bold=True)
    add_bullets(
        doc,
        [
            "Responsive mobile-first UI",
            "Secure API with auth middleware and CORS / Helmet",
            "Reasonable page-load and image optimization (Cloudinary)",
            "Deployable on cloud (Vercel + Railway)",
            "Maintainable modular React and Express code",
        ],
    )
    add_para(doc, "Hardware / Software Environment", bold=True)
    add_bullets(
        doc,
        [
            "Laptop / phone browser; Android device for app testing",
            "Node.js 18+, npm",
            "MongoDB Atlas cluster",
            "Firebase project, Cashfree account, Cloudinary account",
            "VS Code / Cursor IDE, GitHub, Android Studio (for AAB)",
        ],
    )

    add_heading_custom(doc, "3.2 Proposed System Architecture", 2)
    add_para(
        doc,
        "CrwdCtrl follows a classic client–server architecture with external services for auth, "
        "payments and media.",
    )
    add_para(doc, "High-level flow:", bold=True)
    add_para(
        doc,
        "React Client (Web / Capacitor Android)\n"
        "→ REST API on Express (Railway)\n"
        "→ MongoDB Atlas (Mongoose models)\n"
        "+ Firebase Auth / FCM\n"
        "+ Cashfree Payments\n"
        "+ Cloudinary Uploads\n"
        "+ Email (Resend / SMTP)",
    )
    add_para(
        doc,
        "Frontend is hosted on Vercel as a SPA with rewrite to index.html. Backend exposes /api routes "
        "for public content, user auth, organizers and admin. Tokens are sent as Bearer JWT for "
        "authenticated calls. Payment webhooks update booking status on the server.",
    )

    add_heading_custom(doc, "3.3 Data Flow Diagram and System Workflow", 2)
    add_para(doc, "Typical user booking workflow:", bold=True)
    add_para(
        doc,
        "Browse / Search → Open detail page → Register / Login → Fill form → "
        "Cashfree payment (if paid) → Confirmation email → Booking + QR ticket → "
        "Organizer scans QR at venue",
    )
    add_para(doc, "Organizer workflow:", bold=True)
    add_para(
        doc,
        "Login to portal → Select event / trek / run → View dashboard stats → "
        "Manage participants / review payments → Scan QR → Send notifications",
    )
    add_para(doc, "Admin workflow:", bold=True)
    add_para(
        doc,
        "Admin login → Manage listings & homepage sections → Approve organizers → "
        "Monitor registrations / analytics",
    )

    add_heading_custom(doc, "3.4 System Modules Description", 2)
    modules = [
        (
            "3.4.1 Home & Category Discovery Module",
            "Shows hero banners and carousels for ongoing events, happening near you, treks, "
            "sports and run clubs. Uses optimized images and skeletons for smooth mobile UX.",
        ),
        (
            "3.4.2 Auth & Profile Module",
            "Handles email/password and Google sign-in, profile, favourites, bookings and notifications.",
        ),
        (
            "3.4.3 Registration & Payments Module",
            "Dynamic forms per event type, Cashfree checkout, booking records and QR tickets.",
        ),
        (
            "3.4.4 Fest / Competition Module",
            "Fest listing pages, detail views, competitions and fest organizer check-in tools.",
        ),
        (
            "3.4.5 Treks & Community Module",
            "Trek listings, community pages, trek booking and trek-organizer portal "
            "(participants, scan, notify).",
        ),
        (
            "3.4.6 Sports & Run Club Module",
            "Sports category, run clubs, run event details/booking and run-club-organizer portal "
            "with payment-proof review.",
        ),
        (
            "3.4.7 Admin CMS Module",
            "Dashboard for content, sections, coupons, analytics, scanner access and organizer management.",
        ),
        (
            "3.4.8 Mobile / PWA Module",
            "Service worker PWA support and Capacitor Android packaging for Play Store AAB builds.",
        ),
    ]
    for title, body in modules:
        add_heading_custom(doc, title, 3)
        add_para(doc, "Purpose", bold=True)
        add_para(doc, body)
    page_break(doc)

    # -------- CH 4 --------
    add_heading_custom(doc, "METHODOLOGY / IMPLEMENTATION", 1)
    add_heading_custom(doc, "4.1 Registration & Booking Flow Description", 2)
    add_para(
        doc,
        "The methodology of CrwdCtrl explains how discovery becomes a confirmed booking and how "
        "organizers verify attendance.",
    )
    add_para(doc, "Core ideas:", bold=True)
    add_bullets(
        doc,
        [
            "Public content is fetched from Express public APIs",
            "Logged-in actions require JWT / organizer tokens",
            "Paid events create a Cashfree order; webhook confirms success",
            "QR payload encodes booking identity for gate scanners",
            "Organizer scan endpoint marks check-in status",
        ],
    )
    add_para(doc, "Example scoring is not used here; instead business rules are applied:", bold=True)
    add_bullets(
        doc,
        [
            "Capacity / seats remaining reduce after confirmed booking",
            "Pending QR payments can expire after configured TTL",
            "Platform fee (~3%) applied on eligible paid runs / treks",
        ],
    )

    add_heading_custom(doc, "4.2 Tools and Technologies Used", 2)
    add_heading_custom(doc, "4.2.1 React.js", 3)
    add_para(
        doc,
        "React is used to build the interactive SPA. Pages are organized by domain "
        "(home, fests, treks, sports, profile, admin, organizer portals). React Router maps "
        "URLs to lazy-loaded page components.",
    )
    add_para(doc, "Use of React.js in CrwdCtrl:")
    add_bullets(
        doc,
        [
            "Reusable UI components (cards, carousels, headers)",
            "Context providers (auth, dark mode, favourites, notifications)",
            "Hooks for data loading and UX polish (page content loading, header collapse)",
        ],
    )

    add_heading_custom(doc, "4.2.2 Vite + Tailwind CSS", 3)
    add_para(
        doc,
        "Vite provides fast development server and production bundling. Tailwind CSS is used for "
        "responsive utility-based styling aligned to CrwdCtrl’s design tokens.",
    )

    add_heading_custom(doc, "4.2.3 Node.js / Express + MongoDB", 3)
    add_para(
        doc,
        "Express exposes modular routers for public, auth, admin, organizers and payments. "
        "Mongoose schemas model fests, events, bookings, clubs, treks and users. MongoDB Atlas "
        "stores production data.",
    )

    add_heading_custom(doc, "4.2.4 Firebase, Cashfree, Cloudinary", 3)
    add_para(
        doc,
        "Firebase supports Google login and push notifications. Cashfree handles payment checkout. "
        "Cloudinary stores and serves optimized cover images.",
    )

    add_heading_custom(doc, "4.2.5 Capacitor Android", 3)
    add_para(
        doc,
        "Capacitor wraps the production web build into an Android app. Release AAB is signed and "
        "uploaded to Play Console (example version 1.0.11 / versionCode 13).",
    )

    add_heading_custom(doc, "4.2.6 Git and GitHub", 3)
    add_para(
        doc,
        "Source control uses GitHub. Frontend deploy triggers from master push to Vercel; backend "
        "deploys on Railway.",
    )

    add_heading_custom(doc, "4.3 Implementation Details and Workflow", 2)
    add_heading_custom(doc, "4.3.1 Project Folder Structure", 3)
    add_para(
        doc,
        "CrwdCtrl/\n"
        "├── frontend/                 # React SPA + Capacitor Android\n"
        "│   ├── src/App.jsx\n"
        "│   ├── src/app/router/       # public, admin, organizer routes\n"
        "│   ├── src/pages/            # feature pages\n"
        "│   ├── src/components/\n"
        "│   ├── src/services/         # API clients\n"
        "│   └── android/              # native project / AAB\n"
        "├── backend/\n"
        "│   └── src/\n"
        "│       ├── server.js\n"
        "│       ├── routers/\n"
        "│       ├── controllers/\n"
        "│       ├── model/\n"
        "│       └── middleware/\n"
        "├── business/                 # brand & positioning docs\n"
        "├── ARCHITECTURE.md\n"
        "└── README.md",
    )

    add_heading_custom(doc, "4.3.2 Frontend Implementation", 3)
    add_para(
        doc,
        "Step 1: Application loads providers and router.\n"
        "Step 2: Public home fetches fests / home feed and auxiliary sections.\n"
        "Step 3: User opens a detail page and starts registration if authenticated.\n"
        "Step 4: Payment SDK / native gateway completes paid checkout.\n"
        "Step 5: Bookings page shows tickets; organizer scan validates QR.",
    )

    add_heading_custom(doc, "4.3.3 Backend Implementation", 3)
    add_para(
        doc,
        "Controllers validate input, apply auth guards, talk to Mongoose models and return JSON. "
        "Payment webhooks update booking status. Organizer routes verify organizer session tokens. "
        "Admin routes require admin privileges.",
    )

    add_heading_custom(doc, "4.4 Module-wise Development", 2)
    add_para(doc, "PROJECT TIMELINE", bold=True)
    add_table(
        doc,
        ["SR.NO", "ACTIVITY", "HOURS"],
        [
            ["1", "Requirement analysis and topic selection", "4"],
            ["2", "Study of existing discovery / ticketing approaches", "4"],
            ["3", "System design & architecture", "6"],
            ["4", "Frontend hubs, cards, detail pages", "20"],
            ["5", "Backend APIs, models, auth", "18"],
            ["6", "Payments, bookings, QR check-in", "12"],
            ["7", "Organizer + admin portals", "16"],
            ["8", "Android Capacitor packaging & AAB", "6"],
            ["9", "Testing, UX polish, deployment", "10"],
        ],
    )
    add_para(
        doc,
        "Major modules developed: Home Dashboard, Fests, Treks, Sports/Run Clubs, Events, Auth/Profile, "
        "Registrations/Payments, Trek Organizer, Run Club Organizer, Admin Panel, PWA/Android.",
    )

    add_heading_custom(doc, "4.5 Sample Screenshots / Outputs", 2)
    add_para(
        doc,
        "[Insert screenshots here while preparing final print PDF]\n"
        "Suggested captures:\n"
        "1) Home page carousels\n"
        "2) Fest / trek / run detail page\n"
        "3) Registration + payment screen\n"
        "4) Bookings / QR ticket\n"
        "5) Run club / trek organizer dashboard (Dash / Guests / Scan / Notify)\n"
        "6) Admin section manager\n"
        "7) Android app home",
    )
    page_break(doc)

    # -------- CH 5 --------
    add_heading_custom(doc, "EXPERIMENTAL RESULTS, RESULTS AND DISCUSSION", 1)
    add_heading_custom(doc, "5.1 Experimental Setup", 2)
    add_para(
        doc,
        "After completing core features, testing was done to check whether modules work on web and Android.",
    )
    add_para(doc, "Development servers:")
    add_bullets(
        doc,
        [
            "Frontend Vite: http://localhost:5173/",
            "Backend Express: http://localhost:8080/api",
            "Production: https://www.crwdctrl.in with Railway API",
        ],
    )
    add_para(
        doc,
        "Browsers used: Chrome desktop and Android Chrome. Android package tested via signed release AAB.",
    )

    add_heading_custom(doc, "5.2 Testing and Test Cases", 2)
    add_para(doc, "Testing Approach: manual functional testing + API response checks.")
    add_table(
        doc,
        ["ID", "Module", "Input / Action", "Expected", "Result"],
        [
            ["TC01", "Home", "Open /", "Home carousels / skeleton then content", "Pass"],
            ["TC02", "Auth", "Google / email login", "Session created, profile available", "Pass"],
            ["TC03", "Browse", "Open Sports / Treks", "Listings load; revisit uses cache", "Pass"],
            ["TC04", "Register", "Submit booking form", "Creates registration draft / booking", "Pass"],
            ["TC05", "Payment", "Cashfree success flow", "Booking confirmed", "Pass"],
            ["TC06", "QR", "Open ticket QR", "QR renders for check-in", "Pass"],
            ["TC07", "Organizer", "Dash / Scan nav", "Page opens (no HTML download)", "Pass"],
            ["TC08", "Admin", "Update homepage section", "Public home reflects change", "Pass"],
            ["TC09", "Android", "Install AAB build", "App opens production API", "Pass"],
        ],
    )

    add_heading_custom(doc, "5.3 Results and Analysis", 2)
    add_para(
        doc,
        "Functional tests confirm that discovery, auth, booking and organizer flows work in the implemented MVP. "
        "UX fixes (image placeholders, session caches, organizer button navigation) reduced page-open glitches "
        "on mobile.",
    )
    add_para(
        doc,
        "Performance observations: Cloudinary presets and eager-loading for centre/hero images improve perceived "
        "speed. Service worker was adjusted to avoid caching /api responses that caused empty or stale feeds.",
    )

    add_heading_custom(doc, "5.4 Discussion", 2)
    add_para(
        doc,
        "CrwdCtrl demonstrates a practical solution to fragmented campus-event discovery. Compared with "
        "WhatsApp/Forms workflows, participants get a clearer catalogue and booking history, while organizers "
        "get check-in and participant tools. Remaining challenges include growing real supply of events, "
        "organizer onboarding, and continuous mobile polish.",
    )
    page_break(doc)

    # -------- CH 6 --------
    add_heading_custom(doc, "CONCLUSION AND FUTURE WORK", 1)
    add_heading_custom(doc, "6.1 Summary of Findings", 2)
    add_para(
        doc,
        "CrwdCtrl was developed as a full-stack community and event discovery platform. The project "
        "includes user discovery across fests, treks, sports/run clubs and events; registration and "
        "payments; organizer portals; admin CMS; and Android packaging.",
    )

    add_heading_custom(doc, "6.2 Achievements and Contributions", 2)
    add_para(doc, "Major Achievements")
    add_bullets(
        doc,
        [
            "Working multi-category discovery experience on web",
            "Auth + bookings + Cashfree payments path",
            "Role-based organizer dashboards with QR scan",
            "Admin tools for content and homepage sections",
            "Production deployment (Vercel + Railway + MongoDB Atlas)",
            "Capacitor Android release AAB pipeline",
        ],
    )

    add_heading_custom(doc, "6.3 Limitations and Improvements", 2)
    add_bullets(
        doc,
        [
            "Marketplace liquidity depends on active organizers and listings",
            "iOS Capacitator build not shipped yet",
            "Some organizer flows still need deeper automation (refunds, settlements)",
            "Recommendation / personalization is basic",
            "Need larger-scale load testing and analytics dashboards",
        ],
    )

    add_heading_custom(doc, "6.4 Future Scope", 2)
    add_bullets(
        doc,
        [
            "Expand city coverage and organizer acquisition",
            "Smarter recommendations based on interests and location",
            "iOS app release",
            "Deeper analytics for organizers (conversion, attendance)",
            "Improved search and map-based discovery",
            "Automated payouts / settlement reports",
        ],
    )
    page_break(doc)

    # -------- REFERENCES --------
    add_heading_custom(doc, "REFERENCES", 1)
    refs = [
        "React Documentation — https://react.dev/",
        "Vite Documentation — https://vitejs.dev/",
        "Express.js Documentation — https://expressjs.com/",
        "MongoDB Manual — https://www.mongodb.com/docs/",
        "Firebase Authentication — https://firebase.google.com/docs/auth",
        "Cashfree Payments Docs — https://www.cashfree.com/docs/",
        "Capacitor Documentation — https://capacitorjs.com/docs",
        "Tailwind CSS Documentation — https://tailwindcss.com/docs",
        "CrwdCtrl product site — https://www.crwdctrl.in/",
        "Project architecture notes — CrwdCtrl ARCHITECTURE.md / README.md (internal)",
    ]
    for i, r in enumerate(refs, 1):
        add_para(doc, f"[{i}] {r}", size=12, space_after=4)

    add_para(doc, "", size=12)
    add_para(
        doc,
        "Note: Insert real screenshots in Section 4.5 before final PDF print.",
        size=11,
        bold=True,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
